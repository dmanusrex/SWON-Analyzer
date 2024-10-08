# DocGen - https://github.com/dmanusrex/docgen
#
# Copyright (C) 2023 - Darren Richer
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.


""" Recommendations (ODP) Main Screen """

import logging
import os
from datetime import datetime
from threading import Thread
from tkinter import BooleanVar, StringVar, filedialog
from typing import Any

import customtkinter as ctk  # type: ignore
import docx  # type: ignore
import pandas as pd
from docx import Document  # type: ignore
from docxcompose.composer import Composer  # type: ignore
from slugify import slugify

# Appliction Specific Imports
from config import AnalyzerConfig
from CTkMessagebox import CTkMessagebox  # type: ignore
from rtr import RTR
from rtr_fields import RTR_CLINICS
from ui_common import Officials_Status_Frame
from tooltip import ToolTip

tkContainer = Any


class Generate_Documents_Frame(ctk.CTkFrame):
    """Generate Word Documents from a supplied RTR file"""

    def __init__(self, container: tkContainer, config: AnalyzerConfig, rtr: RTR):
        super().__init__(container)
        self._config = config
        self._rtr = rtr

        # Get the needed options
        self._officials_list = StringVar(value=self._config.get_str("officials_list"))
        self._officials_list_filename = StringVar(value=os.path.basename(self._officials_list.get()))
        self._odp_report_directory = StringVar(value=self._config.get_str("odp_report_directory"))
        self._odp_report_file = StringVar(value=self._config.get_str("odp_report_file_docx"))

        # Add support for the club selection
        self._club_list = ["None"]
        self._club_selected = ctk.StringVar(value="None")

        # self is a vertical container that will contain 3 frames
        self.columnconfigure(0, weight=1)

        # Setup the frames

        ctk.CTkLabel(self, text="Officials Recommendations").grid(column=0, row=0, sticky="we", pady=(10, 0))

        club_select_frame = ctk.CTkFrame(self)
        club_select_frame.grid(column=0, row=1, columnspan=2, sticky="news", padx=10, pady=10)

        Officials_Status_Frame(self, self._config).grid(column=0, row=2, sticky="news")

        filesframe = ctk.CTkFrame(self)
        filesframe.grid(column=0, row=3, sticky="news")
        filesframe.rowconfigure(0, weight=1)
        filesframe.rowconfigure(1, weight=1)
        filesframe.rowconfigure(2, weight=1)

        buttonsframe = ctk.CTkFrame(self)
        buttonsframe.grid(column=0, row=4, sticky="news")
        buttonsframe.rowconfigure(0, weight=0)

        # Files Section
        ctk.CTkLabel(filesframe, text="Files and Directories").grid(column=0, row=0, sticky="w", padx=10)

        btn2 = ctk.CTkButton(filesframe, text="Recommendations Folder", command=self._handle_report_dir_browse)
        btn2.grid(column=0, row=1, padx=20, pady=10, sticky="ew")
        ToolTip(btn2, text="Select where output files will be sent")
        ctk.CTkLabel(filesframe, textvariable=self._odp_report_directory).grid(
            column=1, row=1, sticky="w", padx=(0, 10)
        )

        btn3 = ctk.CTkButton(filesframe, text="Consolidated Report File", command=self._handle_report_file_browse)
        btn3.grid(column=0, row=2, padx=20, pady=10, sticky="ew")
        ToolTip(btn3, text="Set report file name")
        ctk.CTkLabel(filesframe, textvariable=self._odp_report_file).grid(column=1, row=2, sticky="w", padx=(0, 10))

        # Right options frame for status options

        self.club_dropdown = ctk.CTkOptionMenu(
            club_select_frame, dynamic_resizing=True, values=self._club_list, variable=self._club_selected
        )
        self.club_dropdown.grid(row=0, column=1, padx=10, pady=(10, 10), sticky="w")
        ctk.CTkLabel(club_select_frame, text="Club", anchor="w").grid(
            row=0, column=0, sticky="w", padx=10, pady=(10, 10)
        )

        # Add Command Buttons

        ctk.CTkLabel(buttonsframe, text="Actions").grid(column=0, row=0, sticky="w", padx=10)

        self.reports_btn = ctk.CTkButton(buttonsframe, text="Generate Reports", command=self._handle_reports_btn)
        self.reports_btn.grid(column=0, row=1, sticky="ew", padx=20, pady=10)

        self.bar = ctk.CTkProgressBar(master=buttonsframe, orientation="horizontal", mode="indeterminate")

        # Register Callback
        self._rtr.register_update_callback(self.refresh_club_list)

    def refresh_club_list(self):
        self._club_list = ["None"]
        self._club_list = self._club_list + [club[1] for club in self._rtr.club_list_names]

        self.club_dropdown.configure(values=self._club_list)
        if len(self._club_list) == 2:  # There is only one club loaded, change default
            self.club_dropdown.set(self._club_list[1])
        else:
            self.club_dropdown.set(self._club_list[0])
        logging.info("Officials Development - Club List Refreshed")

    def _handle_report_dir_browse(self) -> None:
        directory = filedialog.askdirectory()
        if len(directory) == 0:
            return
        directory = os.path.normpath(directory)
        self._config.set_str("odp_report_directory", directory)
        self._odp_report_directory.set(directory)

    def _handle_report_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            title="Report File",
            initialfile=os.path.basename(self._odp_report_file.get()),
            initialdir=self._config.get_str("odp_report_directory"),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("odp_report_file_docx", report_file)
        self._odp_report_file.set(report_file)

    def buttons(self, newstate) -> None:
        """Enable/disable all buttons"""
        self.reports_btn.configure(state=newstate)

    def _handle_reports_btn(self) -> None:
        if self._rtr.rtr_data.empty:
            logging.info("Load data first...")
            CTkMessagebox(master=self, title="Error", message="Load RTR Data First", icon="cancel", corner_radius=0)
            return
        club = self._club_selected.get()
        if club == "None":
            logging.info("Select a club first...")
            CTkMessagebox(master=self, title="Error", message="Select a club first", icon="cancel", corner_radius=0)
            return
        self.buttons("disabled")
        self.bar.grid(row=2, column=0, pady=10, padx=20, sticky="s")
        self.bar.set(0)
        self.bar.start()
        reports_thread = Generate_Reports(self._rtr, self._config, club)
        reports_thread.start()
        self.monitor_reports_thread(reports_thread)

    def monitor_reports_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms
            self.after(100, lambda: self.monitor_reports_thread(thread))
        else:
            self.buttons("enabled")
            self.bar.stop()
            self.bar.grid_forget()
            thread.join()


class docgenCore:
    def __init__(self, club: str, club_data_set: pd.DataFrame, config: AnalyzerConfig, **kwargs):
        self._club_data_full = club_data_set
        self._club_data = self._club_data_full.query("Level < 4")
        self.club_code = club

        self._config = config

    def _get_date(self, date_string) -> str:
        if pd.isnull(date_string):
            return ""
        return date_string

    def add_clinic(self, table: Any, clinic_name: str, entry: Any, pos_info: dict) -> None:
        row = table.add_row().cells
        row[0].text = clinic_name
        row[1].text = self._get_date(entry[pos_info["clinicDate"]])
        if pos_info["deckEvals"]:
            row[2].text = self._get_date(entry[pos_info["deckEvals"]].iloc[0])
            if len(entry[pos_info["deckEvals"]]) > 1:
                row[3].text = self._get_date(entry[pos_info["deckEvals"]].iloc[1])
            else:
                row[3].text = "N/A"
        else:
            row[2].text = "N/A"
            row[3].text = "N/A"

        row[0].width = docx.shared.Inches(2.5)
        row[1].width = docx.shared.Inches(1.5)
        row[2].width = docx.shared.Inches(1.5)
        row[3].width = docx.shared.Inches(1.5)
        row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def dump_data_docx(self, club_fullname: str, reportdate: str) -> list:
        """Produce the Word Document for the club and return a list of files"""

        _report_directory = self._config.get_str("odp_report_directory")
        _email_list_csv = self._config.get_str("email_list_csv")
        csv_list = []  # CSV entries for email list (Lastname, Firstname, E-Mail address and Filename)

        for index, entry in self._club_data.iterrows():
            # create a filename from the last and firstnames using slugify and the report directory

            filename = os.path.abspath(
                os.path.join(_report_directory, slugify(entry["Last Name"] + "_" + entry["First Name"]) + ".docx")
            )

            doc = Document()

            doc.add_heading("2024/25 Officials Development", 0)

            p = doc.add_paragraph()
            p.add_run("Report Date: " + reportdate)
            p.add_run(
                "\n\nName: "
                + entry["Last Name"]
                + ", "
                + entry["First Name"]
                + " (SNC ID # "
                + entry["Registration Id"]
                + ")"
            )
            p.add_run("\n\nClub: " + club_fullname + " (" + self.club_code + ")")
            p.add_run("\n\nCurrent Certification Level: ")
            p.add_run(
                "NONE" if pd.isnull(entry["Current_CertificationLevel"]) else entry["Current_CertificationLevel"]
            )

            table = doc.add_table(rows=1, cols=4)
            row = table.rows[0].cells
            row[0].text = "Clinic"
            row[1].text = "Clinic Date"
            row[2].text = "Sign Off #1"
            row[3].text = "Sign Off #2"
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            self.add_clinic(table, "Intro to Swimming", entry, RTR_CLINICS["Intro"])
            self.add_clinic(table, "Safety Marshal", entry, RTR_CLINICS["Safety"])
            self.add_clinic(table, "Stroke & Turn (Combo)", entry, RTR_CLINICS["ST"])
            self.add_clinic(table, "Inspector of Turns", entry, RTR_CLINICS["IT"])
            self.add_clinic(table, "Judge of Stroke", entry, RTR_CLINICS["JoS"])
            self.add_clinic(table, "Chief Timekeeper", entry, RTR_CLINICS["CT"])
            self.add_clinic(table, "Admin Desk (Clerk)", entry, RTR_CLINICS["AdminDesk"])
            self.add_clinic(table, "Meet Manager", entry, RTR_CLINICS["MM"])
            self.add_clinic(table, "Starter", entry, RTR_CLINICS["Starter"])
            self.add_clinic(table, "CFJ/CJE", entry, RTR_CLINICS["CFJ"])
            self.add_clinic(table, "Chief Recorder/Recorder", entry, RTR_CLINICS["ChiefRec"])
            self.add_clinic(table, "Referee", entry, RTR_CLINICS["Referee"])
            self.add_clinic(table, "Para eModule", entry, RTR_CLINICS["Para"])

            table.style = "Light Grid Accent 5"
            table.autofit = True

            # Add logic to define pathway progression

            doc.add_heading("Recommended Actions", 2)

            Intro_Signoffs = entry["Intro_Count"]
            IT_Signoffs = entry["IT_Count"]
            JoS_Signoffs = entry["JoS_Count"]
            Combo_Signoffs = entry["ST_Count"]

            # For NoLevel officials, add a section to identify what they need to do to get to Level I

            if entry["Level"] == 0:
                if entry["Intro_Status"] == "N":
                    doc.add_paragraph("Take Introduction to Swimming Officiating Clinic", style="List Bullet")
                elif Intro_Signoffs < 2:
                    doc.add_paragraph(
                        f"Obtain {2-Intro_Signoffs} sign-off(s) as a Timer",
                        style="List Bullet",
                    )

                if entry["Safety_Status"] == "N":
                    doc.add_paragraph("Take Safety Marshal Clinc", style="List Bullet")

            # For Level I officials - check if they have stroke & turn and have completed 2 sign-offsj

            # With addition of the separated clinics the recommendation logic has been modified as:
            # Combo clinic - 2 sign offs - Recommend JoS sign-off if they don't have one.
            # Combo clinic - < 2 sign offs - Recommend IT sign-offs + JoS sign-off
            # Split clinic - # of neeed IT sign offs if < 2
            # Split clinic - Take JoS if they don't have it and have at least 1 IT sign-off
            # Split Clinic - JoS Sign-off if still need it
            # If they have most of the sign-offs tell them to take a L2 clinic. 3 out of the 5 needed signoffs.

            if entry["Level"] == 1 and Intro_Signoffs > 0:
                if Intro_Signoffs < 2:
                    doc.add_paragraph(
                        f"Obtain {2-Intro_Signoffs} sign-off(s) as a Timer",
                        style="List Bullet",
                    )
                if entry["ST_Status"] == "N":  # They don't have the combo clinic
                    if entry["IT_Status"] == "N":  # They don't have the new IT clinic either
                        doc.add_paragraph("Take Inspector of Turns Clinic and obtain 2 sign-offs", style="List Bullet")
                    elif Intro_Signoffs < 2:  # They don't have all their Timer sign-offs yet
                        doc.add_paragraph(
                            f"Obtain {2-Intro_Signoffs} sign-off(s) as a Timer",
                            style="List Bullet",
                        )
                    if Intro_Signoffs > 0:  # They have at least 1 Intro Sign-Off
                        if (
                            entry["JoS_Status"] == "N" and entry["IT_Status"] != "N"
                        ):  # They have the new IT clinic but not the JoS clinic
                            doc.add_paragraph("Take Judge of Stroke Clinic", style="List Bullet")
                        elif JoS_Signoffs == 0:
                            doc.add_paragraph("Obtain 1 sign-off as Judge of Stroke", style="List Bullet")
                else:  # Has the Legacy Combo Clinic
                    if Combo_Signoffs < 2:
                        doc.add_paragraph(
                            f"Obtain {2-Combo_Signoffs} sign-off(s) as Inspector of Turns",
                            style="List Bullet",
                        )
                    if (Combo_Signoffs > 0) and (JoS_Signoffs == 0):
                        doc.add_paragraph("Obtain 1 sign-off as Judge of Stroke", style="List Bullet")

                # Determine Level II clinic recommendations.
                if (Intro_Signoffs + Combo_Signoffs + JoS_Signoffs >= 3) or (
                    Intro_Signoffs + IT_Signoffs + JoS_Signoffs >= 3
                ):  # They have completed or nearly completed "core" requirements
                    # Check if they have any other clinics
                    if (
                        entry["CT_Status"] == "N"
                        and entry["Admin_Status"] == "N"
                        and entry["MM_Status"] == "N"
                        and entry["Starter_Status"] == "N"
                        and entry["CFJ_Status"] == "N"
                    ):
                        doc.add_paragraph(
                            "Take a Level II clinic (CT, MM, CFJ/CJE, Admin Desk or Starter) and obtain sign-offs",
                            style="List Bullet",
                        )
                    else:  # They have a clinic - Check if any are fully signed off. If not recommend that.
                        if not (
                            entry["CT_Count"] == 2
                            or entry["Admin_Count"] == 2
                            or entry["MM_Count"] == 2
                            or entry["Starter_Count"] == 2
                            or entry["CFJ_Count"] == 2
                        ):
                            doc.add_paragraph(
                                "Obtain sign-offs on at least 1 Level II clinic (CT, MM, CFJ/CJE, Admin Desk or Starter)",
                                style="List Bullet",
                            )
            elif Intro_Signoffs < 2:
                doc.add_paragraph(
                    f"Obtain {2-Intro_Signoffs} sign-off(s) as a Timer",
                    style="List Bullet",
                )

            # If they are a referee and don't have the Para e-module or Domestic clinic
            #    para_status = RTR_CLINICS["Para"]
            #    pentry = entry["Para Swimming eModule"]
            #                paradom_status = RTR_CLINICS["ParaDom"]["hasClinic"]
            #                if entry["Referee_Status"] != "N" and (
            #                    (entry[para_status] != "yes") or (entry[paradom_status] != "Trained Official")
            #                ):
            #                    doc.add_paragraph("Take the Para-Swimming e-Module")
            try:
                doc.save(filename)
                csv_list.append(
                    [entry["Last Name"], entry["First Name"], entry["Email"], filename]
                )  # Only add if saved
            except Exception as e:
                logging.info(
                    f'Error processing offiical {entry["Last Name"]}, {entry["First Name"]}: {type(e).__name__} - {e}'
                )

        return csv_list


class Generate_Reports(Thread):
    def __init__(self, rtr: RTR, config: AnalyzerConfig, selected_club: str):
        super().__init__()
        self._rtr = rtr
        self._df: pd.DataFrame = self._rtr.rtr_data
        self._config: AnalyzerConfig = config
        self._selected_club = selected_club

    def run(self):
        # This still has code elements to run multiple clubs. For single clubs a simple filter has been applied.
        # Once the final specifications are set, this can be greatly simplified.

        logging.info("Reporting in Progress...")

        _report_directory = self._config.get_str("odp_report_directory")
        _report_file_docx = self._config.get_str("odp_report_file_docx")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_docx))
        _email_list_csv = self._config.get_str("email_list_csv")
        _full_csv_file = os.path.abspath(os.path.join(_report_directory, _email_list_csv))

        club_list_names = self._rtr.club_list_names

        club_summaries = []

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        all_csv_entries = []

        for club, club_full in filter(lambda x: x[1] == self._selected_club, club_list_names):
            logging.info("Processing %s" % club_full)
            club_data = self._df[(self._df["ClubCode"] == club)]
            club_data = club_data[club_data["Status"].isin(status_values)]
            club_stat = docgenCore(club, club_data, self._config)
            club_csv = club_stat.dump_data_docx(club_full, report_time)
            all_csv_entries.extend(club_csv)
            club_summaries.append([club, club_full, club_stat])

        # Create the email list CSV file
        #
        # The email list is a CSV file with the following columns:
        #  Last Name, First Name, E-Mail address, Filename

        logging.info("Creating email list CSV file")
        email_list_df = pd.DataFrame(all_csv_entries, columns=["Last Name", "First Name", "EMail", "Filename"])

        try:
            email_list_df.to_csv(_full_csv_file, index=False)
        except Exception as e:
            logging.info("Unable to save email list: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))

        # Create the master document

        logging.info("Creating master document")

        number_of_sections = len(all_csv_entries)
        master = Document()
        composer = Composer(master)
        for i in range(0, number_of_sections):
            try:
                doc_temp = Document(all_csv_entries[i][3])
            except Exception as e:
                logging.info("Unable to load document: {}".format(type(e).__name__))
                logging.info("Exception message: {}".format(e))
                continue
            doc_temp.add_page_break()
            composer.append(doc_temp)

        try:
            composer.save(_full_report_file)
        except Exception as e:
            logging.info("Unable to save full report: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))

        logging.info("Report Complete")


def main():
    """testing"""
    root = ctk.CTk()
    root.resizable(True, True)
    options = AnalyzerConfig()
    rtrdata = RTR(options)
    settings = Generate_Documents_Frame(root, options, rtrdata)
    settings.grid(column=0, row=0, sticky="news")
    root.mainloop()


if __name__ == "__main__":
    main()
