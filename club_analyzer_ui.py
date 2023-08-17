# Club Analyzer - https://github.com/dmanusrex/SWON-Analyzer
# Copyright (C) 2023 - Darren Richer
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.


''' Analyzer Main Screen '''

import os
import pandas as pd
import numpy as np
import logging
import webbrowser
import tkinter as tk
import customtkinter as ctk
from tkinter import filedialog, ttk, BooleanVar, StringVar
from typing import Any
from threading import Thread
from datetime import datetime
from copy import deepcopy, copy
from tooltip import ToolTip
from docx import Document

# Appliction Specific Imports
from config import AnalyzerConfig
from version import ANALYZER_VERSION
from club_summary import club_summary

tkContainer = Any

class TextHandler(logging.Handler):
    # This class allows you to log to a Tkinter Text or ScrolledText widget
    # Adapted from Moshe Kaplan: https://gist.github.com/moshekaplan/c425f861de7bbf28ef06

    def __init__(self, text):
        # run the regular Handler __init__
        logging.Handler.__init__(self)
        # Store a reference to the Text it will log to
        self.text = text

    def emit(self, record):
        msg = self.format(record)
        def append():
            self.text.configure(state='normal')
            self.text.insert(tk.END, msg + '\n')
            self.text.configure(state='disabled')
            # Autoscroll to the bottom
            self.text.yview(tk.END)
        # This is necessary because we can't modify the Text from other threads
        self.text.after(0, append)

class _Data_Loader(Thread):
    '''Load Data files'''
    def __init__(self, config: AnalyzerConfig):
        super().__init__()
        self._config = config
        self.df : pd.DataFrame 
        self.affiliates : pd.DataFrame 

    def run(self):
        html_file = self._config.get_str("officials_list")
        self.club_list_names_df = pd.DataFrame
        self.club_list_names = []
        logging.info("Loading RTR Data")
        try:
            self.df = pd.read_html(html_file)[0]
        except:
            logging.info("Unable to load data file")
            self.df = pd.DataFrame
            self.affiliates = pd.DataFrame
            return
        self.df.columns = self.df.iloc[0]   # The first row is the column names
        self.df = self.df[1:]

        # Club Level exports include blank rows, purge those out
        self.df.drop(index=self.df[self.df['Registration Id'].isnull()].index, inplace=True)

        # The RTR has 2 types of "empty" dates.  One is blank the other is 0001-01-01.  Fix that.
        self.df.replace('0001-01-01', np.nan, inplace=True)

        # The RTR export is inconsistent on column values for certifications. Fix that.
        self.df.replace('Yes','yes', inplace=True)    # We don't use the no value so no need to fix it 

        logging.info("Loaded %d officials" % self.df.shape[0])

        logging.info("Loading Complete")

class _Generate_Reports(Thread):
    def __init__(self, df: pd.DataFrame, affiliates: pd.DataFrame, config: AnalyzerConfig):
        super().__init__()
        self._df : pd.DataFrame = df
        self._affiliates : pd.DataFrame = affiliates
        self._config : AnalyzerConfig = config

    def run(self):
        logging.info("Reporting in Progress...")

        _report_directory = self._config.get_str("report_directory")
        _report_file_docx = self._config.get_str("report_file_docx")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_docx))
        _full_report = self._config.get_bool("gen_word")
        _per_club = self._config.get_bool("gen_1_per_club")
        _use_affiliates = self._config.get_bool("incl_affiliates")

        club_list_names_df = self._df.loc[self._df['AffiliatedClubs'].isnull(),['ClubCode','Club']].drop_duplicates()
        club_list_names = club_list_names_df.values.tolist()
        club_list_names.sort(key=lambda x:x[0])

        club_summaries = []

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        if _full_report:
            doc = Document()
        for club, club_full in club_list_names:
            logging.info("Processing %s" % club_full)
            affiliation_reg_ids = []
            if _use_affiliates and not self._affiliates.empty:
                affiliation_club_list = self._affiliates[self._affiliates["AffiliatedClubs"] == club]
                if not affiliation_club_list.empty:
                    affiliation_reg_ids = affiliation_club_list[("Registration Id")].values.tolist()
            club_data = self._df[(self._df["ClubCode"] == club) | (self._df["Registration Id"].isin(affiliation_reg_ids))]
            club_data = club_data[club_data["Status"].isin(status_values)]
            club_stat = club_summary(club, club_data, self._config)
#                club_stat.dump_data(f, club_full)
            if _full_report:
                club_stat.dump_data_docx(doc, club_full, report_time, affiliation_reg_ids)
                doc.add_page_break()
            if _per_club:
                _club_file=os.path.abspath(os.path.join(_report_directory, club+".docx"))
                club_doc = Document()
                club_stat.dump_data_docx(club_doc, club_full, report_time, affiliation_reg_ids)
                try:
                    club_doc.save(_club_file)
                except Exception as e:
                    logging.info("Unable to save individual report: {}".format(type(e).__name__))
                    logging.info("Exception message: {}".format(e))
            club_summaries.append ([club, club_full, club_stat])
        
        if _full_report:
            try:
                doc.save(_full_report_file)
            except Exception as e:
                logging.info("Unable to save full report: {}".format(type(e).__name__))
                logging.info("Exception message: {}".format(e))

        logging.info("Reports Complete")

class _Cohost_Analyzer(Thread):
    def __init__(self, df: pd.DataFrame, affiliates: pd.DataFrame, config: AnalyzerConfig, selected_clubs: list):
        super().__init__()
        self._df : pd.DataFrame = df
        self._affiliates : pd.DataFrame = affiliates
        self._config : AnalyzerConfig = config
        self._selected_clubs : list = selected_clubs

    def run(self):
        logging.info("Co-hosting analysis in Progress...")

        _report_directory = self._config.get_str("report_directory")
        _report_file_cohost = self._config.get_str("report_file_cohost")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_cohost))
        _use_affiliates = self._config.get_bool("incl_affiliates")

        club_list_names = self._selected_clubs
        club_codes = [i[0] for i in self._selected_clubs]
        club_full = ""
        comma = False

        for i in club_list_names:
            if comma:
                club_full += ', \n'
            else:
                comma = True
            club_full += i[1]

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%A %B %d %Y %I:%M%p")

        doc = Document()
        logging.info("Processing Co-Hosting for %s" % club_full)
        affiliation_reg_ids = []
        if _use_affiliates and not self._affiliates.empty:
            # Affiliation is removed if the offiical is from one of the host clubs
            affiliation_club_list = self._affiliates[self._affiliates["AffiliatedClubs"].isin(club_codes) & ~self._affiliates["ClubCode"].isin(club_codes)]
            if not affiliation_club_list.empty:
                affiliation_reg_ids = affiliation_club_list[("Registration Id")].values.tolist()
        club_data = self._df[(self._df["ClubCode"].isin(club_codes)) | (self._df["Registration Id"].isin(affiliation_reg_ids))]
        club_data = club_data[club_data["Status"].isin(status_values)]
        club_stat = club_summary("COHOST", club_data, self._config)
        club_stat.dump_data_docx(doc, club_full, report_time, affiliation_reg_ids)
        try:
            doc.save(_full_report_file)
        except:
            logging.info("Unable to save report")

        logging.info("Co-Hosting Analysis Complete")

class _Configuration_Files(ctk.CTkFrame):   # pylint: disable=too-many-ancestors
    '''The "start list" portion of the settings'''
    def __init__(self, container: tkContainer, config: AnalyzerConfig):
        super().__init__(container)
        self._config = config
        self._officials_list = StringVar(value=self._config.get_str("officials_list"))
        self._report_directory = StringVar(value=self._config.get_str("report_directory"))
        self._ctk_theme = StringVar(value=self._config.get_str("Theme"))
        self._ctk_size = StringVar(value=self._config.get_str("Scaling"))
        self._ctk_colour = StringVar(value=self._config.get_str("Colour"))

         # self is a vertical container
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        
        # row 0: label
        lbl1 = ctk.CTkLabel(self, text="Files, Directories and UI Options")
        lbl1.grid(column=0, row=0, sticky="ws")
        # row 1: browse button 
        # fr1 is horizontal
        fr1 = ctk.CTkFrame(self)
        fr1.grid(column=0, row=1, sticky="news")
        fr1.rowconfigure(0, weight=1)
        officials_label = ctk.CTkLabel(fr1, textvariable=self._officials_list)
        officials_label.grid(column=1, row=0, sticky="ew")
        btn1 = ctk.CTkButton(fr1, text="RTR List", command=self._handle_officials_browse)
        btn1.grid(column=0, row=0, padx=20, pady=10)
        ToolTip(btn1, text="Select the RTR officials export file")   # pylint: disable=C0330
        # row 2: reserved for multi-datafile loader

        # row 3: Output Directory
        fr3 = ctk.CTkFrame(self)
        fr3.grid(column=0, row=3, sticky="news")
        fr3.rowconfigure(0, weight=1)
        output_directory_label = ctk.CTkLabel(fr3, textvariable=self._report_directory)
        output_directory_label.grid(column=1, row=0, sticky="ew")
        btn3 = ctk.CTkButton(fr3, text="Report Directory", command=self._handle_report_dir_browse)
        btn3.grid(column=0, row=0, padx=20, pady=10)
        ToolTip(btn3, text="Select where output files will be sent")   # pylint: disable=C0330

        fr4 = ctk.CTkFrame(self)
        fr4.grid(column=0, row=4, sticky="news")
        fr4.rowconfigure(0, weight=1)
        self.appearance_mode_label = ctk.CTkLabel(fr4, text="Appearance Mode", anchor="w")
        self.appearance_mode_label.grid(row=0, column=1)
        self.appearance_mode_optionemenu = ctk.CTkOptionMenu(fr4, values=["Light", "Dark", "System"],
                                                                       command=self.change_appearance_mode_event, variable=self._ctk_theme)
        self.appearance_mode_optionemenu.grid(row=0, column=0, padx=20, pady=10)

        fr5 = ctk.CTkFrame(self)
        fr5.grid(column=0, row=5, sticky="news")
        fr5.rowconfigure(0, weight=1)
        self.scaling_label = ctk.CTkLabel(fr5, text="UI Scaling", anchor="w")
        self.scaling_label.grid(row=0, column=1)
        self.scaling_optionemenu = ctk.CTkOptionMenu(fr5, values=["80%", "90%", "100%", "110%", "120%"],
                                                               command=self.change_scaling_event, variable=self._ctk_size)
        self.scaling_optionemenu.grid(row=0, column=0, padx=20, pady=10)
  

        fr6 = ctk.CTkFrame(self)
        fr6.grid(column=0, row=6, sticky="news")
        fr6.rowconfigure(0, weight=1)
        self.colour_label = ctk.CTkLabel(fr6, text="Colour (Application Restart Required)", anchor="w")
        self.colour_label.grid(row=0, column=1)
        self.colour_optionemenu = ctk.CTkOptionMenu(fr6, values=["blue", "green", "dark-blue"],
                                                               command=self.change_colour_event, variable=self._ctk_colour)
        self.colour_optionemenu.grid(row=0, column=0, padx=20, pady=10)

    def _handle_officials_browse(self) -> None:
        directory = filedialog.askopenfilename()
        if len(directory) == 0:
            return
        self._config.set_str("officials_list", directory)
        self._officials_list.set(directory)

    def _handle_report_dir_browse(self) -> None:
        directory = filedialog.askdirectory()
        if len(directory) == 0:
            return
        directory = os.path.normpath(directory)
        self._config.set_str("report_directory", directory)
        self._report_directory.set(directory)

    def change_appearance_mode_event(self, new_appearance_mode: str):
        ctk.set_appearance_mode(new_appearance_mode)
        self._config.set_str("Theme", new_appearance_mode)

    def change_scaling_event(self, new_scaling: str):
        new_scaling_float = int(new_scaling.replace("%", "")) / 100
        ctk.set_widget_scaling(new_scaling_float)
        self._config.set_str("Scaling", new_scaling)

    def change_colour_event(self, new_colour: str):
        logging.info("Changing colour to : " + new_colour)
        ctk.set_default_color_theme(new_colour)
        self._config.set_str("Colour", new_colour)

class _AnalyzerTab(ctk.CTkFrame):  # pylint: disable=too-many-ancestors,too-many-instance-attributes
    '''Miscellaneous settings'''
    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container) # , text="General Settings", padding=5
        self._config = config
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=0)

        self.rowconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)
        self.rowconfigure(3, weight=1)
        self.rowconfigure(4, weight=1)
        self.rowconfigure(5, weight=1)

        # Row 0 - Labels
        lbl1 = ctk.CTkLabel(self, text="Status Options")
        lbl1.grid(column=0, row=0, sticky="ws", pady=10)

        lbl2 = ctk.CTkLabel(self, text="Analyzer/Report Settings")
        lbl2.grid(column=1, row=0, sticky="ws", pady=10)

        # Row 1 - Account Pending, Include RTR Errors
        self._incl_account_pending().grid(column=0, row=1, sticky="ws")
        self._incl_errors().grid(column=1, row=1, sticky="ws")

        # Row 2 - PSO Pending, Sanction Errors
        self._incl_pso_pending().grid(column=0, row=2, sticky="ws")
        self._incl_sanction_errors().grid(column=1, row=2, sticky="ws")

        # Row 3 - Invoice Pending, Use Affiliates
        self._incl_inv_pending().grid(column=0, row=3, sticky="ws")
        self._incl_affiliates().grid(column=1, row=3, sticky="ws")

    def _incl_errors(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_errors_var = ctk.BooleanVar(frame, value=self._config.get_bool("incl_errors"))
        ctk.CTkSwitch(frame, text="RTR Errors/Warnings", command=self._handle_incl_errors, variable=self._incl_errors_var,
                      onvalue = True, offvalue=False).grid(column=0, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include errors and anomalies detected")
        return frame

    def _handle_incl_errors(self, *_arg):
        self._config.set_bool("incl_errors", self._incl_errors_var.get())

    def _incl_inv_pending(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_inv_pending_var = BooleanVar(frame, value=self._config.get_bool("incl_inv_pending"))
        ctk.CTkSwitch(frame, text = "Invoice Pending", variable=self._incl_inv_pending_var, onvalue = True, offvalue=False,
            command=self._handle_incl_inv_pending).grid(column=1, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include Invoice Pending status")
        return frame
    
    def _handle_incl_inv_pending(self, *_arg):
        self._config.set_bool("incl_inv_pending", self._incl_inv_pending_var.get())

    def _incl_pso_pending(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_pso_pending_var = BooleanVar(frame, value=self._config.get_bool("incl_pso_pending"))
        ctk.CTkSwitch(frame, text = "PSO Pending", variable=self._incl_pso_pending_var, onvalue = True, offvalue=False,
            command=self._handle_incl_pso_pending).grid(column=1, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include PSO Pending status")
        return frame
    
    def _handle_incl_pso_pending(self, *_arg):
        self._config.set_bool("incl_pso_pending", self._incl_pso_pending_var.get())

    def _incl_account_pending(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_account_pending_var = BooleanVar(frame, value=self._config.get_bool("incl_account_pending"))
        ctk.CTkSwitch(frame, text = "Account Pending", variable=self._incl_account_pending_var, onvalue = True, offvalue=False,
            command=self._handle_incl_account_pending).grid(column=1, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include Account Pending status")
        return frame
    
    def _handle_incl_account_pending(self, *_arg):
        self._config.set_bool("incl_account_pending", self._incl_account_pending_var.get())


    def _incl_affiliates(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_affiliates_var = BooleanVar(frame, value=self._config.get_bool("incl_affiliates"))
        ctk.CTkSwitch(frame, text = "Affiliates", variable=self._incl_affiliates_var, onvalue = True, offvalue=False,
            command=self._handle_incl_affiliates).grid(column=1, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include affiliated officials in the analysis")
        return frame
    
    def _handle_incl_affiliates(self, *_arg):
        self._config.set_bool("incl_affiliates", self._incl_affiliates_var.get())

    def _incl_sanction_errors(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._incl_sanction_errors_var = ctk.BooleanVar(frame, value=self._config.get_bool("incl_sanction_errors"))
        ctk.CTkSwitch(frame, text="Sanction Errors", command=self._handle_incl_sanction_errors, variable=self._incl_sanction_errors_var,
                      onvalue = True, offvalue=False).grid(column=0, row=0, sticky="n", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Select to include why the sanctioning tier could not be permitted")
        return frame

    def _handle_incl_sanction_errors(self, *_arg):
        self._config.set_bool("incl_sanction_errors", self._incl_sanction_errors_var.get())

class _ReportsTab(ctk.CTkFrame):  # pylint: disable=too-many-ancestors,too-many-instance-attributes
    '''Reports Setting and Generation'''
    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container) # , text="General Settings", padding=5
        self._config = config
        self._report_file = StringVar(value=self._config.get_str("report_file_docx"))
        self._report_directory = StringVar(value=self._config.get_str("report_directory"))

        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)
        self.columnconfigure(2, weight=1)
        self.rowconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)

        # Row 0 - Labels
        lbl3 = ctk.CTkLabel(self, text="Report Options")
        lbl3.grid(column=1, row=0, sticky="ws", pady=10)


        # Row 1 - Account Pending, Include Errors
        self._gen_1_per_club().grid(column=1, row=1, sticky="ws")

        # Row 2 - PSO Pending, 1 per club
        self._gen_word().grid(column=1, row=2, sticky="ws")

        # Row 3 - Report File
        report_file_label = ctk.CTkLabel(self, textvariable=self._report_file)
        report_file_label.grid(column=1, row=4, sticky="ew", padx=2, pady=(0,10))
        btn1 = ctk.CTkButton(self, text="Master Report File Name", command=self._handle_report_file_browse)
        btn1.grid(column=1, row=3, padx=20, pady=(20, 10))


    def _gen_1_per_club(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._gen_1_per_club_var = BooleanVar(frame, value=self._config.get_bool("gen_1_per_club"))
        ctk.CTkSwitch(frame, text="Individual Files", variable=self._gen_1_per_club_var, onvalue = True, offvalue=False,
            command=self._handle_gen_1_per_club).grid(column=1, row=0, sticky="news", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Generate 1 document per club")
        return frame

    def _handle_gen_1_per_club(self, *_arg):
        self._config.set_bool("gen_1_per_club", self._gen_1_per_club_var.get())

    def _gen_word(self) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(self)
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self._gen_word_var = BooleanVar(frame, value=self._config.get_bool("gen_word"))
        ctk.CTkSwitch(frame, text="Main Report", variable=self._gen_word_var, onvalue = True, offvalue=False,
            command=self._handle_gen_word).grid(column=1, row=0, sticky="news", padx=20, pady=10) # pylint: disable=C0330
        ToolTip(frame, "Generate the master word document")
        return frame

    def _handle_gen_word(self, *_arg):
        self._config.set_bool("gen_word", self._gen_word_var.get())

    def _handle_report_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename( filetypes = [('Word Documents','*.docx')], defaultextension=".docx", title="Report File", 
                                                initialfile=os.path.basename(self._report_file.get()),
                                                initialdir=self._config.get_str("report_directory"))
        if len(report_file) == 0:
            return
        self._config.set_str("report_file_docx", report_file)
        self._report_file.set(report_file)

class _CohostTab(ctk.CTkFrame):  # pylint: disable=too-many-ancestors,too-many-instance-attributes
    '''Co-Hosting'''
    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container) # , text="General Settings", padding=5
        self._config = config
        self._club_list = ['None']
        self._club_list_full = []   
        self._report_file = StringVar(value=self._config.get_str("report_file_cohost"))
        self._report_directory = StringVar(value=self._config.get_str("report_directory"))

        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)
        self.club1 = ctk.StringVar(value='None')
        self.club2 = ctk.StringVar(value='None')
        self.club3 = ctk.StringVar(value='None')
        # Row 0 - Labels
        lbl3 = ctk.CTkLabel(self, text="Report Options")
        lbl3.grid(column=0, row=0, sticky="ws", pady=10)


        self.club1_dropdown = ctk.CTkOptionMenu(self, dynamic_resizing=True,
                                                        values=self._club_list, variable=self.club1)
        self.club1_dropdown.grid(row=1, column=0, padx=20, pady=(20, 10))
        self.club2_dropdown = ctk.CTkOptionMenu(self, dynamic_resizing=True,
                                                        values=self._club_list, variable=self.club2)
        self.club2_dropdown.grid(row=2, column=0, padx=20, pady=(20, 10))
        self.club3_dropdown = ctk.CTkOptionMenu(self, dynamic_resizing=True,
                                                        values=self._club_list, variable=self.club3)
        self.club3_dropdown.grid(row=3, column=0, padx=20, pady=(20, 10))

        self.report_file_label = ctk.CTkLabel(self, textvariable=self._report_file)
        self.report_file_label.grid(column=0, row=5, sticky="ew", padx=20, pady=(0, 10))
        btn1 = ctk.CTkButton(self, text="Co-Hosting Report File Name", command=self._handle_cohost_file_browse)
        btn1.grid(column=0, row=4, padx=20, pady=(20, 10))


    def refresh_club_list(self, clublist: list):
        self._club_list_full = clublist
        self._clublist = ['None']
        self._clublist = self._clublist + [club[1] for club in clublist]

        self.club1_dropdown.configure(values=self._clublist)
        self.club2_dropdown.configure(values=self._clublist)
        self.club3_dropdown.configure(values=self._clublist)
        logging.info("Club List Refreshed")

    def get_clubs(self) -> list:
        selected_items = [self.club1.get(), self.club2.get(), self.club3.get()]
        final_list = [i for i in self._club_list_full if i[1] in selected_items]
        return final_list

    def _handle_cohost_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename( filetypes = [('Word Documents','*.docx')], defaultextension=".docx", title="Co-Host Report File", 
                                                initialfile=os.path.basename(self._report_file.get()),
                                                initialdir=self._config.get_str("report_directory"))
        if len(report_file) == 0:
            return
        self._config.set_str("report_file_cohost", report_file)
        self._report_file.set(report_file)

class _Logging(ctk.CTkFrame): # pylint: disable=too-many-ancestors,too-many-instance-attributes
    '''Logging Window'''
    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container) # , text="Logging Window"
        self._config = config
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        lbl1 = ctk.CTkLabel(self, text="Messages")
        lbl1.grid(column=0, row=0, sticky="ws", pady=10)

        self.logwin = ctk.CTkTextbox(self, state='disabled')
        self.logwin.grid(column=0, row=4, sticky='nsew')
        # Logging configuration
        logfile = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "swon-analyzer.log"))

        logging.basicConfig(filename=logfile,
                            level=logging.INFO,
                            format='%(asctime)s - %(levelname)s - %(message)s')
        # Create textLogger
        text_handler = TextHandler(self.logwin)
        text_handler.setFormatter(logging.Formatter('%(levelname)s - %(message)s'))
        # Add the handler to logger
        logger = logging.getLogger()
        logger.addHandler(text_handler)

class SwonApp(ctk.CTkFrame):  # pylint: disable=too-many-ancestors
    '''Main Appliction'''
    # pylint: disable=too-many-arguments,too-many-locals
    def __init__(self, container: ctk.CTk,
                 config: AnalyzerConfig):
        super().__init__(container)
        self._config = config
        self.df = pd.DataFrame()
        self.affiliates = pd.DataFrame()
        self.grid(column=0, row=0, sticky="news")
        self.columnconfigure(0, weight=1)
        # Odd rows are empty filler to distribute vertical whitespace
        for i in [1, 3, 5, 7]:
            self.rowconfigure(i, weight=1)

        # Use Tabs for readabilty

        self.tabview = ctk.CTkTabview(self, width=container.winfo_width())
        self.tabview.grid(row=0, column=0, padx=(20, 0), pady=(20, 0), sticky="nsew")
        self.tabview.add("Configuration")
        self.tabview.add("Analyzer/Report Settings")
        self.tabview.add("Reports")
        self.tabview.add("Co-Hosting")

        # Configuration Tab
        self.tabview.tab("Configuration").grid_columnconfigure(0, weight=1)
        configfiles = _Configuration_Files(self.tabview.tab("Configuration"), self._config)
        configfiles.grid(column=0, row=0, sticky="news")
        self.load_btn = ctk.CTkButton(self.tabview.tab("Configuration"), text="Load Datafile", command=self._handle_load_btn)
        self.load_btn.grid(column=0, row=1, sticky="news", padx=20, pady=10)
        self.reset_btn = ctk.CTkButton(self.tabview.tab("Configuration"), text="Reset", command=self._handle_reset_btn)
        self.reset_btn.grid(column=0, row=2, sticky="news", padx=20, pady=10)


        # Analyzer Tab
        self.tabview.tab("Analyzer/Report Settings").grid_columnconfigure(0, weight=1)
        analyzer = _AnalyzerTab(self.tabview.tab("Analyzer/Report Settings"), self._config)
        analyzer.grid(column=0, row=0, sticky="news")

        # Report Tab
        self.tabview.tab("Reports").grid_columnconfigure(0, weight=1)
        reports = _ReportsTab(self.tabview.tab("Reports"), self._config)
        reports.grid(column=0, row=0, sticky="")
        self.reports_btn = ctk.CTkButton(self.tabview.tab("Reports"), text="Generate Reports", command=self._handle_reports_btn)
        self.reports_btn.grid(column=0, row=1, sticky="")

        # Co-Hosting Tab
        self.tabview.tab("Co-Hosting").grid_columnconfigure(0, weight=1)
        self.cohost = _CohostTab(self.tabview.tab("Co-Hosting"), self._config)
        self.cohost.grid(column=0, row=0, sticky="")
        self.cohost_btn = ctk.CTkButton(self.tabview.tab("Co-Hosting"), text="Co-Hosting Report", command=self._handle_cohost_btn)
        self.cohost_btn.grid(column=0, row=1, sticky="", padx=20, pady=10)


        # Logging Window
        loggingwin = _Logging(self, self._config)
        loggingwin.grid(column=0, row=4, sticky="news")
 
        # Info panel
        fr8 = ctk.CTkFrame(self)
        fr8.grid(column=0, row=8, sticky="news")
        fr8.rowconfigure(0, weight=1)
        fr8.columnconfigure(0, weight=1)
        link_label = ctk.CTkLabel(fr8,
            text="Documentation: https://swon-analyzer.readthedocs.io")  # pylint: disable=C0330
        link_label.grid(column=0, row=0, sticky="w")
        # Custom Tkinter clickable label example https://github.com/TomSchimansky/CustomTkinter/issues/1208
        link_label.bind("<Button-1>", lambda event: webbrowser.open("https://swon-analyzer.readthedocs.io")) # link the command function
        link_label.bind("<Enter>", lambda event: link_label.configure(font=("",13,"underline"), cursor="hand2"))
        link_label.bind("<Leave>", lambda event: link_label.configure(font=("",13), cursor="arrow"))

        version_label = ctk.CTkLabel(fr8, text="Version "+ANALYZER_VERSION)
        version_label.grid(column=1, row=0, sticky="nes")

    def buttons(self, newstate) -> None:
        '''Enable/disable all buttons on the UI'''
        self.load_btn.configure(state = newstate)
        self.reset_btn.configure(state = newstate)
        self.reports_btn.configure(state = newstate)
        self.cohost_btn.configure(state = newstate)

    def _handle_reports_btn(self) -> None:
        if self.df.empty:
            logging.info ("Load data first...")
            return
        self.buttons("disabled")
        reports_thread = _Generate_Reports(self.df, self.affiliates, self._config)
        reports_thread.start()
        self.monitor_reports_thread(reports_thread)


    def _handle_load_btn(self) -> None:
        self.buttons("disabled")
        load_thread = _Data_Loader(self._config)
        load_thread.start()
        self.monitor_load_thread(load_thread)

    def _handle_reset_btn(self) -> None:
        self.buttons("disabled")
        self.df = pd.DataFrame()
        self.affiliates = pd.DataFrame()
        self.club_list_names_df = pd.DataFrame()
        self.club_list_names = []
        self.cohost.refresh_club_list(self.club_list_names)
        logging.info("Reset Complete")
        self.buttons("enabled")


    def _handle_cohost_btn(self) -> None:
        if self.df.empty:
            logging.info ("Load data first...")
            return
        self.buttons("disabled")
        club_list = self.cohost.get_clubs()

        if club_list:
            cohost_thread = _Cohost_Analyzer(self.df, self.affiliates, self._config, club_list)
            cohost_thread.start()
            self.monitor_cohost_thread(cohost_thread)
        else:
            logging.info("Please select at least 1 club first")
            self.buttons("enabled")

    def monitor_load_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms 
            self.after(100, lambda: self.monitor_load_thread(thread))
        else:
            # Retrieve data from the loading process and merge it with already loaded data
            if self.df.empty:
                self.df = thread.df
            else:
                self.df = pd.concat([self.df,thread.df], axis=0).drop_duplicates()
                logging.info("%d officials records merged" % self.df.shape[0])
 
            # We exclude affiliated offiicals from determining the list of clubs. This is important for club level exports.
            self.club_list_names_df = self.df.loc[self.df['AffiliatedClubs'].isnull(),['ClubCode','Club']].drop_duplicates()
            self.club_list_names = self.club_list_names_df.values.tolist()
            self.club_list_names.sort(key=lambda x:x[0])

            logging.info("Extracting Affiliation Data")

            # Find officials with affiliated clubs. For sanctioning affiliated offiicals must have a certification level
            self.affiliates = self.df[~self.df['AffiliatedClubs'].isnull() & ~self.df['Current_CertificationLevel'].isnull()].copy()
            if self.affiliates.empty:
                logging.info("No affiliation records found")
            else:
                self.affiliates['AffiliatedClubs'] = self.affiliates['AffiliatedClubs'].str.split(',')
                self.affiliates = self.affiliates.explode('AffiliatedClubs')
                logging.info("Extracted %d affiliation records" % self.affiliates.shape[0])

            self.cohost.refresh_club_list(self.club_list_names)
            self.buttons("enabled")
            thread.join()

    def monitor_reports_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms 
            self.after(100, lambda: self.monitor_reports_thread(thread))
        else:
            self.buttons("enabled")
            thread.join()
        
    def monitor_cohost_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms 
            self.after(100, lambda: self.monitor_cohost_thread(thread))
        else:
            self.buttons("enabled")
            thread.join()
        
def main():
    '''testing'''
    root = ctk.CTk()
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    root.resizable(True, True)
    options = AnalyzerConfig()
    settings = SwonApp(root, options)
    settings.grid(column=0, row=0, sticky="news")
    logging.info("Hello World")
    root.mainloop()

if __name__ == '__main__':
    main()
