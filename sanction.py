# Club Analyzer - https://github.com/dmanusrex/SWON-Analyzer
# Copyright (C) 2023 - Darren Richer
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.


""" Sanctioning Application """

import os
import pandas as pd
import logging
import customtkinter as ctk  # type: ignore
from tkinter import filedialog, BooleanVar, StringVar
from threading import Thread
from datetime import datetime
from tooltip import ToolTip
from docx import Document  # type: ignore
from CTkMessagebox import CTkMessagebox  # type: ignore


# Appliction Specific Imports
from config import AnalyzerConfig
from club_summary import club_summary
from rtr import RTR
from ui_common import Officials_Status_Frame


class Sanctioning_Report_Options_Frame(ctk.CTkFrame):
    """Sanctioning Report Options Options"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container)
        self._config = config

        self._incl_errors_var = ctk.BooleanVar(value=self._config.get_bool("incl_errors"))
        self._incl_affiliates_var = BooleanVar(value=self._config.get_bool("incl_affiliates"))
        self._incl_sanction_errors_var = ctk.BooleanVar(value=self._config.get_bool("incl_sanction_errors"))

        # self is a vertical container that will contain 1 frame
        self.columnconfigure(0, weight=1)

        # Options Frame - Left and Right Panels

        optionsframe = self
#        optionsframe.grid(column=0, row=0, sticky="news", padx=10, pady=10)
#        optionsframe.rowconfigure(0, weight=1)

        ctk.CTkLabel(optionsframe, text="Report Settings").grid(column=0, row=0, sticky="w", padx=10)

        ctk.CTkSwitch(
            optionsframe,
            text="RTR Errors/Warnings",
            command=self._handle_incl_errors,
            variable=self._incl_errors_var,
            onvalue=True,
            offvalue=False,
        ).grid(column=0, row=2, sticky="w", padx=20, pady=10)

        ctk.CTkSwitch(
            optionsframe,
            text="Affiliated Officials",
            variable=self._incl_affiliates_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_incl_affiliates,
        ).grid(column=0, row=3, sticky="w", padx=20, pady=10)

        ctk.CTkSwitch(
            optionsframe,
            text="Sanction Errors",
            command=self._handle_incl_sanction_errors,
            variable=self._incl_sanction_errors_var,
            onvalue=True,
            offvalue=False,
        ).grid(column=0, row=4, sticky="w", padx=20, pady=10)

    def _handle_incl_affiliates(self, *_arg):
        self._config.set_bool("incl_affiliates", self._incl_affiliates_var.get())

    def _handle_incl_errors(self, *_arg):
        self._config.set_bool("incl_errors", self._incl_errors_var.get())

    def _handle_incl_sanction_errors(self, *_arg):
        self._config.set_bool("incl_sanction_errors", self._incl_sanction_errors_var.get())


class Sanctioning_Options_Frame(ctk.CTkFrame):
    """Sanctioning Options"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig):
        super().__init__(container)
        self._config = config

        self._contractor_results_var = BooleanVar(value=self._config.get_bool("contractor_results"))
        self._contractor_mm_var = BooleanVar(value=self._config.get_bool("contractor_mm"))
        self._video_finish_var = BooleanVar(value=self._config.get_bool("video_finish"))

        # self is a vertical container that will contain 1 frame
        self.columnconfigure(0, weight=1)

        # Options Frame - Left and Right Panels

        optionsframe = self
#        optionsframe.grid(column=0, row=0, sticky="news", padx=10, pady=10)
#        optionsframe.rowconfigure(0, weight=1)

        ctk.CTkLabel(optionsframe, text="Sanctioning Options").grid(column=0, row=0, sticky="w", padx=10)

        ctk.CTkSwitch(
            optionsframe,
            text="Contractor for Results",
            variable=self._contractor_results_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_contractor_results,
        ).grid(column=0, row=2, sticky="w", padx=20, pady=10)

        ctk.CTkSwitch(
            optionsframe,
            text="Contractor for Meet Management",
            variable=self._contractor_mm_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_contractor_mm,
        ).grid(column=0, row=3, sticky="w", padx=20, pady=10)

        ctk.CTkSwitch(
            optionsframe,
            text="Video Finish System",
            variable=self._video_finish_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_video_finish,
        ).grid(column=0, row=4, sticky="w", padx=20, pady=10)

    def _handle_contractor_results(self, *_arg):
        self._config.set_bool("contractor_results", self._contractor_results_var.get())

    def _handle_contractor_mm(self, *_arg):
        self._config.set_bool("contractor_mm", self._contractor_mm_var.get())

    def _handle_video_finish(self, *_arg):
        self._config.set_bool("video_finish", self._video_finish_var.get())


class Sanction_ROR(ctk.CTkFrame):
    """Reports Setting and Generation for RORs and POAs"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig, rtr: RTR):
        super().__init__(container)  # , text="General Settings", padding=5
        self._config = config
        self._rtr = rtr
        self._report_file = StringVar(value=self._config.get_str("report_file_docx"))
        self._report_directory = StringVar(value=self._config.get_str("report_directory"))

        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)

        # Setup the sub-frames
        ctk.CTkLabel(self, text="ROR/POA Reporting (Multi-Club)").grid(
            column=0, row=0, sticky="we", pady=10, columnspan=3
        )

        genoptionsframe = ctk.CTkFrame(self)
        genoptionsframe.grid(column=0, row=1, sticky="news", padx=10, pady=10, columnspan=3)
        genoptionsframe.rowconfigure(0, weight=1)

        Officials_Status_Frame(genoptionsframe, self._config).grid(column=0, row=1, sticky="news")
        Sanctioning_Options_Frame(genoptionsframe, self._config).grid(column=1, row=1, sticky="news")
        Sanctioning_Report_Options_Frame(genoptionsframe, self._config).grid(column=2, row=1, sticky="news")

        optionsframe = ctk.CTkFrame(self)
        optionsframe.grid(column=0, row=2, sticky="news", padx=10, pady=10)

        filesframe = ctk.CTkFrame(self)
        filesframe.grid(column=1, row=2, sticky="news", padx=10, pady=10, columnspan=2)
        filesframe.rowconfigure(0, weight=1)
        filesframe.rowconfigure(1, weight=1)
        filesframe.rowconfigure(2, weight=1)

        buttonsframe = ctk.CTkFrame(self)
        buttonsframe.grid(column=0, row=3, sticky="news", padx=10, pady=10, columnspan=3)
        buttonsframe.rowconfigure(0, weight=1)
        buttonsframe.columnconfigure(0, weight=1)

        # Report Options

        self._gen_word_var = BooleanVar(optionsframe, value=self._config.get_bool("gen_word"))
        ctk.CTkSwitch(
            optionsframe,
            text="Main Report",
            variable=self._gen_word_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_gen_word,
        ).grid(column=1, row=0, sticky="news", padx=20, pady=10)

        self._gen_1_per_club_var = BooleanVar(optionsframe, value=self._config.get_bool("gen_1_per_club"))
        ctk.CTkSwitch(
            optionsframe,
            text="Individual Files",
            variable=self._gen_1_per_club_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_gen_1_per_club,
        ).grid(column=1, row=2, sticky="news", padx=20, pady=10)

        # Report File

        rptbtn = ctk.CTkButton(filesframe, text="Main Report File Name", command=self._handle_report_file_browse)
        rptbtn.grid(column=0, row=0, padx=20, pady=10)
        ToolTip(rptbtn, text="Set report file name")
        ctk.CTkLabel(filesframe, textvariable=self._report_file).grid(column=1, row=0, sticky="w")

        rptdirbtn = ctk.CTkButton(filesframe, text="Report Folder", command=self._handle_report_dir_browse)
        rptdirbtn.grid(column=0, row=2, padx=20, pady=10)
        ToolTip(rptdirbtn, text="Select where output files will be sent")
        ctk.CTkLabel(filesframe, textvariable=self._report_directory).grid(column=1, row=2, sticky="w")

        # Action Button

        self.reports_btn = ctk.CTkButton(buttonsframe, text="Generate Reports", command=self._handle_reports_btn)
        self.reports_btn.grid(column=0, row=0, sticky="ew", padx=20, pady=10)

        self.bar = ctk.CTkProgressBar(master=buttonsframe, orientation="horizontal", mode="indeterminate")

    def buttons(self, newstate) -> None:
        """Enable/disable all buttons on the UI"""
        self.reports_btn.configure(state=newstate)

    def _handle_gen_1_per_club(self, *_arg):
        self._config.set_bool("gen_1_per_club", self._gen_1_per_club_var.get())

    def _handle_gen_word(self, *_arg):
        self._config.set_bool("gen_word", self._gen_word_var.get())

    def _handle_report_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            title="Report File",
            initialfile=os.path.basename(self._report_file.get()),
            initialdir=self._config.get_str("report_directory"),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("report_file_docx", report_file)
        self._report_file.set(report_file)

    def _handle_report_dir_browse(self) -> None:
        directory = filedialog.askdirectory()
        if len(directory) == 0:
            return
        directory = os.path.normpath(directory)
        self._config.set_str("report_directory", directory)
        self._report_directory.set(directory)

    def _handle_reports_btn(self) -> None:
        if self._rtr.rtr_data.empty:
            logging.info("Load data first...")
            CTkMessagebox(master=self, title="Error", message="Load RTR Data First", icon="cancel", corner_radius=0)
            return
        self.buttons("disabled")
        self.bar.grid(row=2, column=0, pady=10, padx=20, sticky="s")
        self.bar.set(0)
        self.bar.start()
        reports_thread = _Generate_Reports(self._rtr, self._config)
        reports_thread.start()
        self.monitor_reports_thread(reports_thread)

    def monitor_reports_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms
            self.after(100, lambda: self.monitor_reports_thread(thread))
        else:
            self.bar.stop()
            self.bar.grid_forget()
            self.buttons("enabled")
            thread.join()


class Sanction_COA_CoHost(ctk.CTkFrame):
    """Co-Hosting"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig, rtr: RTR):
        super().__init__(container)
        self._config = config
        self._rtr = rtr
        self._club_list = ["None"]
        self._report_file = StringVar(value=self._config.get_str("report_file_cohost"))
        self._report_directory = StringVar(value=self._config.get_str("report_directory"))

        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        self.club1 = ctk.StringVar(value="None")
        self.club2 = ctk.StringVar(value="None")
        self.club3 = ctk.StringVar(value="None")

        # Setup the sub-frames
        ctk.CTkLabel(self, text="Sanctioning Reporting").grid(column=0, row=0, sticky="we", pady=10, columnspan=3)

        optionsframe = ctk.CTkFrame(self)
        optionsframe.grid(column=0, row=1, sticky="news", padx=10, pady=10)
        optionsframe.rowconfigure(0, weight=1)

        Officials_Status_Frame(optionsframe, self._config).grid(column=0, row=1, sticky="news")
        Sanctioning_Options_Frame(optionsframe, self._config).grid(column=1, row=1, sticky="news")
        Sanctioning_Report_Options_Frame(optionsframe, self._config).grid(column=2, row=1, sticky="news")

        dropdownframe = ctk.CTkFrame(self)
        dropdownframe.grid(column=0, row=2, sticky="news", padx=10, pady=10, columnspan=3)
        dropdownframe.rowconfigure(0, weight=1)

        filesframe = ctk.CTkFrame(self)
        filesframe.grid(column=0, row=4, sticky="news", padx=10, pady=10, columnspan=3)
        filesframe.rowconfigure(0, weight=1)

        buttonsframe = ctk.CTkFrame(self)
        buttonsframe.grid(column=0, row=6, sticky="news", padx=10, pady=10, columnspan=3)
        buttonsframe.rowconfigure(0, weight=1)
        buttonsframe.columnconfigure(0, weight=1)

        # Club Dropdowns

        ctk.CTkLabel(dropdownframe, text="Select Club(s)").grid(column=0, row=0, sticky="w", padx=15)

        self.club1_dropdown = ctk.CTkOptionMenu(
            dropdownframe, dynamic_resizing=True, values=self._club_list, variable=self.club1
        )
        self.club1_dropdown.grid(row=1, column=0, padx=20, pady=(20, 10), sticky="w")

        self.club2_dropdown = ctk.CTkOptionMenu(
            dropdownframe, dynamic_resizing=True, values=self._club_list, variable=self.club2
        )

        self.club3_dropdown = ctk.CTkOptionMenu(
            dropdownframe, dynamic_resizing=True, values=self._club_list, variable=self.club3
        )

        # Report File

        rptbtn = ctk.CTkButton(filesframe, text="Sanctioning Report Filename", command=self._handle_cohost_file_browse)
        rptbtn.grid(column=0, row=0, padx=20, pady=(20, 10))
        ToolTip(rptbtn, text="Set report file name")
        ctk.CTkLabel(filesframe, textvariable=self._report_file).grid(column=1, row=0, sticky="w", pady=5)

        # Action Button

        self.cohost_btn = ctk.CTkButton(buttonsframe, text="Sanctioning Report", command=self._handle_cohost_btn)
        self.cohost_btn.grid(column=0, row=0, sticky="ew", padx=20, pady=10)

        self._rtr.register_update_callback(self.refresh_club_list)

    def refresh_club_list(self):
        self._club_list = ["None"]
        self._club_list = self._club_list + [club[1] for club in self._rtr.club_list_names]

        self.club1_dropdown.configure(values=self._club_list)
        if len(self._club_list) == 2:  # There is only one club loaded, change default
            self.club1_dropdown.set(self._club_list[1])
        else:
            self.club1_dropdown.set(self._club_list[0])
        self.club2_dropdown.configure(values=self._club_list)
        self.club3_dropdown.configure(values=self._club_list)

        if len(self._club_list) > 2:
            self.club2_dropdown.grid(row=2, column=0, padx=20, pady=(20, 10), sticky="w")
            self.club3_dropdown.grid(row=3, column=0, padx=20, pady=(20, 10), sticky="w")
        else:
            self.club2_dropdown.grid_forget()
            self.club3_dropdown.grid_forget()

        logging.info("Sanctioning - Club List Refreshed")

    def get_clubs(self) -> list:
        selected_items = [self.club1.get(), self.club2.get(), self.club3.get()]
        final_list = [i for i in self._rtr.club_list_names if i[1] in selected_items]
        return final_list

    def buttons(self, newstate) -> None:
        """Enable/disable all buttons on the UI"""
        self.cohost_btn.configure(state=newstate)

    def _handle_cohost_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            title="Co-Host Report File",
            initialfile=os.path.basename(self._report_file.get()),
            initialdir=self._config.get_str("report_directory"),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("report_file_cohost", report_file)
        self._report_file.set(report_file)

    def _handle_cohost_btn(self) -> None:
        if self._rtr.rtr_data.empty:
            logging.info("Load data first...")
            CTkMessagebox(master=self, title="Error", message="Load RTR Data First", icon="cancel", corner_radius=0)
            return
        self.buttons("disabled")
        club_list = self.get_clubs()

        if club_list:
            cohost_thread = _Cohost_Analyzer(self._rtr, self._config, club_list)
            cohost_thread.start()
            self.monitor_cohost_thread(cohost_thread)
        else:
            logging.info("Please select at least 1 club first")
            CTkMessagebox(
                master=self,
                title="Error",
                message="Please select at least 1 club first",
                icon="cancel",
                corner_radius=0,
            )
            self.buttons("enabled")

    def monitor_cohost_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms
            self.after(100, lambda: self.monitor_cohost_thread(thread))
        else:
            self.buttons("enabled")
            thread.join()


class _Generate_Reports(Thread):
    def __init__(self, rtr: RTR, config: AnalyzerConfig):
        super().__init__()
        self._df: pd.DataFrame = rtr.rtr_data
        self._affiliates: pd.DataFrame = rtr.affiliates
        self._club_list_names_df: pd.DataFrame = rtr.club_list_names_df
        self._club_list_names: list = rtr.club_list_names
        self._config: AnalyzerConfig = config

    def run(self):
        logging.info("Reporting in Progress...")

        _report_directory = self._config.get_str("report_directory")
        _report_file_docx = self._config.get_str("report_file_docx")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_docx))
        _full_report = self._config.get_bool("gen_word")
        _per_club = self._config.get_bool("gen_1_per_club")
        _use_affiliates = self._config.get_bool("incl_affiliates")

        club_summaries = []

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        if _full_report:
            doc = Document()
        for club, club_full in self._club_list_names:
            logging.info("Processing %s" % club_full)
            affiliation_reg_ids = []
            if _use_affiliates and not self._affiliates.empty:
                affiliation_club_list = self._affiliates[self._affiliates["AffiliatedClubs"] == club]
                if not affiliation_club_list.empty:
                    affiliation_reg_ids = affiliation_club_list[("Registration Id")].values.tolist()

            club_data = self._df[
                (self._df["ClubCode"] == club) | (self._df["Registration Id"].isin(affiliation_reg_ids))
            ]
            club_data = club_data[club_data["Status"].isin(status_values)]
            club_stat = club_summary(club, club_data, self._config)

            if _full_report:
                club_stat.dump_data_docx(doc, club_full, report_time, affiliation_reg_ids)
                doc.add_page_break()
            if _per_club:
                _club_file = os.path.abspath(os.path.join(_report_directory, club + ".docx"))
                club_doc = Document()
                club_stat.dump_data_docx(club_doc, club_full, report_time, affiliation_reg_ids)
                try:
                    club_doc.save(_club_file)
                except Exception as e:
                    logging.info("Unable to save individual report: {}".format(type(e).__name__))
                    logging.info("Exception message: {}".format(e))
                    CTkMessagebox(
                        title="Error", message=f"Unable to save file for {club_full}", icon="cancel", corner_radius=0
                    )
            club_summaries.append([club, club_full, club_stat])

        if _full_report:
            try:
                doc.save(_full_report_file)
            except Exception as e:
                logging.info("Unable to save full report: {}".format(type(e).__name__))
                logging.info("Exception message: {}".format(e))
                CTkMessagebox(title="Error", message="Unable to save full report file", icon="cancel", corner_radius=0)

        CTkMessagebox(title="Reports", message="Reports complete", icon="check", option_1="OK", corner_radius=0)

        logging.info("Reports Complete")


class _Cohost_Analyzer(Thread):
    def __init__(self, rtr: RTR, config: AnalyzerConfig, selected_clubs: list):
        super().__init__()
        # quick fix - just map new datastructure to old
        self._df: pd.DataFrame = rtr.rtr_data
        self._affiliates: pd.DataFrame = rtr.affiliates
        self._club_list_names_df: pd.DataFrame = rtr.club_list_names_df
        self._club_list_names: list = rtr.club_list_names
        self._config: AnalyzerConfig = config
        self._selected_clubs: list = selected_clubs

    def run(self):
        logging.info("Sanctioning report in Progress...")

        _report_directory = self._config.get_str("report_directory")
        _report_file_cohost = self._config.get_str("report_file_cohost")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_cohost))
        _use_affiliates = self._config.get_bool("incl_affiliates")

        club_list_names = self._selected_clubs
        club_codes = [i[0] for i in self._selected_clubs]
        club_full = "\n".join([item[1] for item in club_list_names])

        if len(club_codes) == 1:
            report_club_code = club_codes[0]
        else:
            report_club_code = "COHOST"

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%A %B %d %Y %I:%M%p")

        doc = Document()
        logging.info("Processing COA/Co-host report for %s" % club_full)

        affiliation_reg_ids = []
        if _use_affiliates and not self._affiliates.empty:
            # Affiliation is removed if the offiical is from one of the host clubs
            affiliation_club_list = self._affiliates[
                self._affiliates["AffiliatedClubs"].isin(club_codes) & ~self._affiliates["ClubCode"].isin(club_codes)
            ]
            if not affiliation_club_list.empty:
                affiliation_reg_ids = affiliation_club_list[("Registration Id")].values.tolist()

        club_data = self._df[
            (self._df["ClubCode"].isin(club_codes)) | (self._df["Registration Id"].isin(affiliation_reg_ids))
        ]
        club_data = club_data[club_data["Status"].isin(status_values)]
        club_stat = club_summary(report_club_code, club_data, self._config)

        club_stat.dump_data_docx(doc, club_full, report_time, affiliation_reg_ids)

        try:
            doc.save(_full_report_file)
            CTkMessagebox(
                title="Sanctioning Report", message="Report complete", icon="check", option_1="OK", corner_radius=0
            )
        except:
            CTkMessagebox(title="Error", message="Unable to save report file", icon="cancel", corner_radius=0)
            logging.info("Unable to save report")

        logging.info("COA/Co-Hosting Analysis Complete")


def main():
    """testing"""
    root = ctk.CTk()
    root.resizable(True, True)
    options = AnalyzerConfig()
    rtrdata = RTR(options)
    # settings = Sanctioning_Report_Options_Frame(root, options)
    # settings = Officials_Status_Frame(root, options)
    # settings = Sanction_ROR(root, options, rtrdata)
    settings = Sanction_COA_CoHost(root, options, rtrdata)
    settings.grid(column=0, row=0, sticky="news")
    root.mainloop()


if __name__ == "__main__":
    main()
