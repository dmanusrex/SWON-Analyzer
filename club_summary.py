# Club Analyzer - https://github.com/dmanusrex/SWON-Analyzer
# Copyright (C) 2023 - Darren Richer
#
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.

""" Club Summary

Main object that does the heavy lifting.  This will produce the required dataset for a given dataframe.
The dataframe does not have to be an individual club. It can be a combo to assess compliances as well.

General Notes:

The RTR data exported for officials is not a complete data set.  Further, the underlying data has a number
of inconsistencies for historical reasons and bad data validation.   The officials export is a normalization of
a number of SQL server tables.

The principle goal is to allow for automated analysis to support a number of activities:

1)  Basic statistical information for clubs
2)  Identify data errors for offiicials
3)  Identify "easy wins" for officials to advances for club
4)  Determine what level of meet sanction a club can apply for.

This class performs all of the necessary functions to create the data for those objectives.

TO DO: Re-factor code to be cleaner. The current version is the result of a number of spec changes over a few
  days and is messy

"""

import logging
from copy import copy, deepcopy
from datetime import datetime
from typing import Any

import docx  # type: ignore
import pandas as pd
from docx import Document  # type: ignore
from docx.shared import Inches  # type: ignore

from config import AnalyzerConfig
from rtr_fields import RTR_CLINICS

LIST_OR_DICT = list | dict


class club_summary:
    def __init__(self, club: str, club_data_set: pd.DataFrame, config: AnalyzerConfig, **kwargs):
        self._club_data_full = club_data_set.copy()
        self._club_data = self._club_data_full.query("Level < 4")
        self.club_code = club
        self._config = config

        # These just store the counts of officials at each level
        self.Level_None: int = 0
        self.Level_1s: int = 0
        self.Level_2s: int = 0
        self.Level_3s: int = 0
        self.Level_4s: int = 0
        self.Level_5s: int = 0
        self.Qual_Refs: int = 0

        # Each list contains a summary count of the number of offiicals with 0, 1, or 2 certification dates
        self.Intro: list = []
        self.SandT: list = []
        self.IT: list = []  # Starting sept/23 IT & Judge of Stroke will be separated
        self.JoS: list = []  # Starting sept/23 IT & Judge of Stroke will be separated
        self.ChiefT: list = []
        self.Clerk: list = []
        self.MM: list = []
        self.Starter: list = []
        self.CFJ: list = []
        self.RecSec: list = []
        self.Referee: list = []
        self.Qualified_Refs: list = []
        self.Level_4_5s: list = []
        self.Level_3_list: list = []

        # These lists contain the names of officials with various RTR errors
        self.NoLevel_Missing_Cert: list = []
        self.NoLevel_Missing_SM: list = []
        self.NoLevel_Has_II: list = []
        self.Missing_Level_II: list = []
        self.Missing_Level_III: list = []

        # List of approved and failed sanctions. Failed sanctions include the failure reason.
        self.Sanction_Level: list = []
        self.Failed_Sanctions: list = []

        # Enable Debug

        self.debug = False

        # Track the best scenario found for each sanctioning level

        self.best_scenario: list = []

        # Build the summary data and check sanctioning abilities
        self._count_levels()
        self._count_certifications()
        self._find_all_level4_5s()
        self._find_qualfied_refs()
        self._check_sanctions()

        # RTR Error Checks
        self._check_no_levels()
        self._check_missing_Level_III()
        self._check_missing_Level_II()

    def _is_valid_date(self, date_string) -> bool:
        if pd.isnull(date_string):
            return False

        try:
            datetime.strptime(date_string, "%Y-%m-%d")
            return True
        except ValueError:
            return False

    def _count_levels(self):
        """Level Statistics"""

        level_counts = self._club_data_full["Level"].value_counts()
        self.Level_None = level_counts.get(0, 0)
        self.Level_1s = level_counts.get(1, 0)
        self.Level_2s = level_counts.get(2, 0)
        self.Level_3s = level_counts.get(3, 0)
        self.Level_4s = level_counts.get(4, 0)
        self.Level_5s = level_counts.get(5, 0)

    def _find_qualfied_refs(self) -> None:
        # To be a Level III referee you need CT, Clerk, Starter and one of CFJ or MM
        # Also check domestic clinic status
        level3_list = self._club_data.query("Level == 3")
        self.Qualified_Refs = []
        self.Level_3_list = []

        for index, row in level3_list.iterrows():
            self.Level_3_list.append(row["Full Name"])
            if (
                row["Referee_Status"] != "N"
                and row["CT_Status"] == "C"
                and row["Admin_Status"] == "C"
                and row["Starter_Status"] == "C"
                and (row["CFJ_Status"] == "C" or row["MM_Status"] == "C")
            ):
                self.Qualified_Refs.append([row["Full Name"], row["Para Domestic"], row["Para Swimming eModule"]])
        self.Qual_Refs = len(self.Qualified_Refs)

    def _find_all_level4_5s(self) -> None:
        """In the RTR Level 4/5s may not have the underlying detail but
        by definition they must be certified in all positions"""

        level45_list = self._club_data_full.query("Level > 3")
        self.Level_4_5s = []

        # Get all their names

        for index, row in level45_list.iterrows():
            self.Level_4_5s.append(row["Full Name"])

    def _check_no_levels(self) -> None:
        no_level_list = self._club_data.query("Level == 0")
        has_both = []
        has_intro_only = []
        has_level_ii = []

        for index, row in no_level_list.iterrows():
            official_name = row["Last Name"] + ", " + row["First Name"]
            if row["Intro_Status"] != "N":
                if row["Safety_Status"] != "N":
                    has_both.append(official_name)  # Missing Certification Record
                else:
                    has_intro_only.append(official_name)  # Missing Safety Marshal
            if (
                (row["ST_Status"] != "N" or (row["IT_Status"] != "N" and row["JoS_Status"] != "N"))
                and row["CT_Status"] != "N"
                and row["Admin_Status"] != "N"
                and row["MM_Status"] != "N"
                and row["Starter_Status"] != "N"
                and row["CFJ_Status"] != "N"
            ):  # Has Level II clinics but missing Level I
                has_level_ii.append(official_name)

        self.NoLevel_Missing_Cert = has_both
        self.NoLevel_Missing_SM = has_intro_only
        self.NoLevel_Has_II = has_level_ii

    def _check_missing_Level_III(self) -> None:
        level_2_list = self._club_data.query("Level == 2")
        self.Missing_Level_III = []
        clinics_to_check = ["CT_Status", "Admin_Status", "Starter_Status", "CFJ_Status", "MM_Status"]

        if level_2_list.empty:
            return

        for index, row in level_2_list.iterrows():
            if (
                row["CT_Status"] == "C"
                and row["Admin_Status"] == "C"
                and row["Starter_Status"] == "C"
                and row["CFJ_Status"] == "C"
                and row["MM_Status"] == "C"
            ):
                cert_count = 0
                for clinic in clinics_to_check:
                    if row[clinic] == "C":
                        cert_count += 1
                if cert_count >= 4:
                    self.Missing_Level_III.append(row["Full Name"])

    def _check_missing_Level_II(self) -> None:
        level_1_list = self._club_data.query("Level == 1")
        self.Missing_Level_II = []
        clinics_to_check = ["CT_Status", "Admin_Status", "Starter_Status", "CFJ_Status", "MM_Status"]

        if level_1_list.empty:
            return

        for index, row in level_1_list.iterrows():
            if row["Intro_Status"] == "C" and (
                row["ST_Status"] == "C" or (row["IT_Status"] == "C" and row["JoS_Status"] == "C")
            ):
                cert_count = 0
                for clinic in clinics_to_check:
                    if row[clinic] == "C":
                        cert_count += 1

                if cert_count >= 1:
                    self.Missing_Level_II.append(row["Full Name"])

    def _count_certifications_detail(self, clinic: dict) -> list:
        cert_total = 0  # Count of clinics taken
        cert_1so = 0  # Has 1 Sign-Off
        cert_2so = 0  # Has 2 Sign-Offs (fully qualified)

        cert_name = clinic["status"]
        cert_count = clinic["signoffs"]

        # Create the certification lists using the officials full name

        cert_list = self._club_data_full.loc[self._club_data_full[cert_name] == "C", ["Full Name"]][
            "Full Name"
        ].values.tolist()
        qual_list = self._club_data_full.loc[self._club_data_full[cert_name] != "N", ["Full Name"]][
            "Full Name"
        ].values.tolist()

        cert_counts = self._club_data_full.loc[self._club_data_full[cert_name] != "N", [cert_count]][
            cert_count
        ].value_counts()
        cert_1so = cert_counts.get(1, 0)
        cert_2so = cert_counts.get(2, 0)
        cert_total = cert_counts.get(0, 0) + cert_1so + cert_2so

        return [cert_total, cert_1so, cert_2so, qual_list, cert_list]

    def _count_certifications(self) -> None:
        self.Intro = self._count_certifications_detail(RTR_CLINICS["Intro"])
        self.SandT = self._count_certifications_detail(RTR_CLINICS["ST"])
        self.IT = self._count_certifications_detail(RTR_CLINICS["IT"])
        self.JoS = self._count_certifications_detail(RTR_CLINICS["JoS"])
        self.ChiefT = self._count_certifications_detail(RTR_CLINICS["CT"])
        self.Clerk = self._count_certifications_detail(RTR_CLINICS["AdminDesk"])
        self.MM = self._count_certifications_detail(RTR_CLINICS["MM"])
        self.Starter = self._count_certifications_detail(RTR_CLINICS["Starter"])
        self.CFJ = self._count_certifications_detail(RTR_CLINICS["CFJ"])
        self.RecSec = self._count_certifications_detail(RTR_CLINICS["ChiefRec"])
        self.Referee = self._count_certifications_detail(RTR_CLINICS["Referee"])

        # For IT and JoS extend their lists to include S&T and drop duplicates (temp fix)
        self.IT[3].extend(self.SandT[3])
        self.IT[4].extend(self.SandT[4])
        self.JoS[3].extend(self.SandT[3])
        self.JoS[4].extend(self.SandT[4])

        self.IT[3] = list(set(self.IT[3]))
        self.IT[4] = list(set(self.IT[4]))
        self.JoS[3] = list(set(self.JoS[3]))
        self.JoS[4] = list(set(self.JoS[4]))

    def _check_sanctions(self) -> None:
        approved_sanctions = []

        """Sanctioning Parameters to pass:

        # of Level 4/5s needed
        # of Level 3 Referees Need
        # of Level 3s needed
        # of CTs needed (qualified, certified)
        # of MMs needed (qualified, certified)
        # of Clerks needed (qualified, certified)
        # of Starters needed (qualified, certified)
        # of CFJs needed (qualified, certified)
        # of ITs needed (qualified, certified)
        # of Judges of Stroke needed (qualified, certified)

        """

        T1Opt1 = self._check_sanctions_detail(1, 0, 0, 1, 0, 1, 0, 0, 0, 1, 0, 1, 0, 4, 0, 0, 0, "TIER I - A")
        T1Opt2 = self._check_sanctions_detail(0, 1, 0, 1, 0, 0, 1, 0, 0, 1, 0, 1, 0, 3, 1, 0, 0, "TIER I - B")

        if T1Opt1 or T1Opt2:
            opts = (
                "TIER I - Class II Time Trial + In-House Competition (Option(s):"
                + (" A" if T1Opt1 else "")
                + (" B" if T1Opt2 else "")
                + ")"
            )
            approved_sanctions.append(opts)

        T2Opt1 = self._check_sanctions_detail(1, 1, 0, 2, 0, 1, 0, 1, 0, 1, 0, 1, 0, 4, 2, 0, 0, "TIER II - A")
        T2Opt2 = self._check_sanctions_detail(1, 0, 1, 1, 1, 1, 0, 1, 0, 1, 0, 1, 0, 4, 2, 0, 0, "TIER II - B")
        T2Opt3 = self._check_sanctions_detail(0, 0, 1, 1, 1, 0, 1, 1, 0, 1, 0, 1, 0, 4, 2, 0, 0, "TIER II - C")

        if T2Opt1 or T2Opt2 or T2Opt3:
            opts = (
                "TIER II - Closed Invitational (limited to 4 sessions) (Option(s): "
                + (" A" if T2Opt1 else "")
                + (" B" if T2Opt2 else "")
                + (" C" if T2Opt3 else "")
                + ")"
            )
            approved_sanctions.append(opts)

        T3Opt1 = self._check_sanctions_detail(1, 1, 1, 2, 0, 1, 0, 1, 0, 1, 0, 1, 0, 6, 2, 1, 0, "TIER III - A")
        T3Opt2 = self._check_sanctions_detail(1, 0, 1, 1, 1, 0, 1, 1, 0, 1, 1, 1, 0, 6, 2, 1, 0, "TIER III - B")

        if T3Opt1 or T3Opt2:
            opts = (
                "TIER III - Open Invitational (limited to 6 sessions, no standards) (Option(s):"
                + (" A" if T3Opt1 else "")
                + (" B" if T3Opt2 else "")
                + ")"
            )
            approved_sanctions.append(opts)

        T4Opt1 = self._check_sanctions_detail(2, 0, 1, 1, 1, 0, 1, 1, 1, 1, 1, 1, 1, 8, 4, 2, 0, "TIER IV - A")
        T4Opt2 = self._check_sanctions_detail(1, 1, 2, 0, 2, 0, 1, 1, 1, 1, 1, 1, 1, 8, 4, 2, 0, "TIER IV - B")

        if T4Opt1 or T4Opt2:
            opts = (
                "TIER IV - Open or Closed Invitational + Regionals & Provincials "
                + "(no session limits, any double-ended meet) (Option(s):"
                + (" A" if T4Opt1 else "")
                + (" B" if T4Opt2 else "")
                + ")"
            )
            approved_sanctions.append(opts)

        self.Sanction_Level = approved_sanctions

    def _check_sanctions_detail(
        self,
        Level4_5: int,
        Qual_Ref: int,
        Level3: int,
        Qual_CT: int,
        Cert_CT: int,
        Qual_MM: int,
        Cert_MM: int,
        Qual_Clerk: int,
        Cert_Clerk: int,
        Qual_Starter: int,
        Cert_Starter: int,
        Qual_CFJ: int,
        Cert_CFJ: int,
        Qual_IT: int,
        Cert_IT: int,
        Qual_JoS: int,
        Cert_JoS: int,
        dbg_scenario_name,
    ) -> dict:
        """build and test sanction application - returns a dictionary of a valid staffing result if found"""

        if self._config.get_bool("contractor_results") and not self._config.get_bool("video_finish"):
            logging.info("Contractor Results Enabled - Skipping Sanctioning Check for CFJ/CJE")
            Qual_CFJ = 0
            Cert_CFJ = 0

        if self._config.get_bool("contractor_mm"):
            logging.info("Contractor Meet Manager Enabled - Downgrading Certified MM to Qualified MM")
            Qual_MM = Qual_MM + Cert_MM
            Cert_MM = 0

        if self._config.get_bool("video_finish"):
            logging.info("Video Finish Enabled - Removing CT Requirement and adding 1 CFJ/CJE")
            Qual_CT = 0
            Cert_CT = 0
            Qual_CFJ += 1

        my_scenario = self._build_staffing_scenario(
            Level4_5,
            Qual_Ref,
            Level3,
            Qual_CT,
            Cert_CT,
            Qual_MM,
            Cert_MM,
            Qual_Clerk,
            Cert_Clerk,
            Qual_Starter,
            Cert_Starter,
            Qual_CFJ,
            Cert_CFJ,
            Qual_IT,
            Cert_IT,
            Qual_JoS,
            Cert_JoS,
        )

        if self.debug:
            logging.debug(self.club_code + ": " + dbg_scenario_name + " - " + str(my_scenario))

        if isinstance(my_scenario, dict):
            self.best_scenario = []
            staff_list = self._find_staffing_scenario(my_scenario, [], len(my_scenario))
            if staff_list:  # Passed Sr. Checks - Check S&T then continue
                if self.debug:
                    logging.debug(self.club_code + ": " + dbg_scenario_name + " - " + str(staff_list))
                SandT_scenario = self._build_staffing_scenario_SandT(Qual_IT, Cert_IT, Qual_JoS, Cert_JoS, staff_list)
                if isinstance(SandT_scenario, dict):
                    if self.debug:
                        logging.debug(self.club_code + ": " + dbg_scenario_name + " - " + str(SandT_scenario))
                    self.best_scenario = []
                    SandT_list = self._find_staffing_scenario(SandT_scenario, [], len(SandT_scenario))
                    if SandT_list:
                        if self.debug:
                            logging.debug(self.club_code + ": " + dbg_scenario_name + " - " + str(SandT_list))
                        staff_list.extend(SandT_list)
                        staff_list.reverse()
                        staff_jobs = list(my_scenario.keys())
                        staff_jobs.extend(list(SandT_scenario.keys()))
                        return dict(zip(staff_jobs, staff_list))
                    else:
                        msg = dbg_scenario_name + " : Unable to staff stroke & turn"
                        if self.debug:
                            logging.debug(self.club_code + ": " + msg)
                        self.Failed_Sanctions.append(msg)
                        return {}
                else:
                    msg = dbg_scenario_name + " : Insufficient remaining stroke & turn"
                    if self.debug:
                        logging.debug(self.club_code + ": " + msg)
                    self.Failed_Sanctions.append(msg)
                    self.Failed_Sanctions.extend(SandT_scenario)
            else:
                msg = dbg_scenario_name + " : Unable to staff senior grid"
                if self.debug:
                    logging.debug(self.club_code + ": " + msg)
                self.Failed_Sanctions.append(msg)
                self.best_scenario.append("FAILED TO STAFF")  # The last job is the one that failed
                self.best_scenario.reverse()
                failed_scenario_jobs = list(my_scenario.keys())[-len(self.best_scenario) :]
                # combine the two lists element-wise into text entries in a list
                failed_scenario = [
                    "  " + failed_scenario_jobs[i] + ": " + self.best_scenario[i]
                    for i in range(len(failed_scenario_jobs))
                ]
                self.Failed_Sanctions.extend(failed_scenario)
                return {}
        else:
            msg = dbg_scenario_name + " : Minimum available skills not met"
            if self.debug:
                logging.debug(self.club_code + ": " + msg)
            self.Failed_Sanctions.append(msg)
            self.Failed_Sanctions.extend(my_scenario)
        return {}

    def _build_staffing_scenario(
        self,
        Level4_5: int,
        Qual_Ref: int,
        Level3: int,
        Qual_CT: int,
        Cert_CT: int,
        Qual_MM: int,
        Cert_MM: int,
        Qual_Clerk: int,
        Cert_Clerk: int,
        Qual_Starter: int,
        Cert_Starter: int,
        Qual_CFJ: int,
        Cert_CFJ: int,
        Qual_IT: int,
        Cert_IT: int,
        Qual_JoS: int,
        Cert_JoS: int,
    ) -> LIST_OR_DICT:
        """Build the dictionary for the senior grid scenario"""

        # Higher Level Positions can backfill lower level positions to achieve the desired outcome

        scenario = {}
        failure_reasons: list = []

        # No need to performance optimize this, we also do some pre-checks to avoid analysis on scenarios
        # we know would fail

        # Check Quick Failure Conditions
        if (
            len(self.ChiefT[3]) < Qual_CT + Cert_CT
            or len(self.ChiefT[4]) < Cert_CT
            or len(self.MM[3]) < Qual_MM + Cert_MM
            or len(self.MM[4]) < Cert_MM
            or len(self.Clerk[3]) < Qual_Clerk + Cert_Clerk
            or len(self.Clerk[4]) < Cert_Clerk
            or len(self.Starter[3]) < Qual_Starter + Cert_Starter
            or len(self.Starter[4]) < Cert_Starter
            or len(self.CFJ[3]) < Qual_CFJ + Cert_CFJ
            or len(self.CFJ[4]) < Cert_CFJ
            or len(self.IT[3]) < Qual_IT
            or len(self.IT[4]) < Cert_IT
            or len(self.JoS[3]) < Qual_JoS
            or len(self.JoS[4]) < Cert_JoS
            or Level4_5 > (self.Level_4s + self.Level_5s)
            or (Qual_Ref + Level4_5) > (self.Level_4s + self.Level_5s + self.Qual_Refs)
            or (Level3 + Level4_5 + Qual_Ref) > (self.Level_5s + self.Level_4s + self.Level_3s)
        ):
            # Return each condition that failed as a list showing required and available
            # For Level 4/5, Level 3 Refs, and Level 3s we need to look at the entire pool when generating the
            # failure message. Calculate the remaining offiicals at each level and use that to generate the message
            # and ensure the remainder is not negative

            # Level 4/5s left
            err_45_level = self.Level_4s + self.Level_5s - Level4_5
            err_45_level = err_45_level if err_45_level > 0 else 0

            # of Level 3 Refs used
            err_3_ref_unfilled = Qual_Ref - (self.Qual_Refs + err_45_level)
            err_3_ref_unfilled = err_3_ref_unfilled if err_3_ref_unfilled > 0 else 0
            err_3_ref_used = Qual_Ref - err_3_ref_unfilled

            # of Level 3s left to staff - Level 3s are a superset of Level 3 refs
            err_3 = Level3 - (self.Level_3s + err_45_level - err_3_ref_used)
            err_3 = err_3 if err_3 > 0 else 0
            err_3_used = Level3 - err_3

            if Level4_5 > (self.Level_4s + self.Level_5s):  # Not enough Level 4/5s
                failure_reasons.append("  Level 4/5: " + str(self.Level_4s + self.Level_5s) + "/" + str(Level4_5))

            if ((Qual_Ref + Level4_5) > (self.Level_4s + self.Level_5s + self.Qual_Refs)) and (Qual_Ref > 0):
                failure_reasons.append("  Level 3 Refs: " + str(err_3_ref_used) + "/" + str(Qual_Ref))
            if ((Level3 + Level4_5 + Qual_Ref) > (self.Level_5s + self.Level_4s + self.Level_3s)) and (Level3 > 0):
                failure_reasons.append("  Level 3s: " + str(err_3_used) + "/" + str(Level3))

            if len(self.ChiefT[3]) < Qual_CT + Cert_CT:
                failure_reasons.append("  CT (Qualified): " + str(len(self.ChiefT[3])) + "/" + str(Qual_CT + Cert_CT))
            if len(self.ChiefT[4]) < Cert_CT:
                failure_reasons.append("  CT (Certified): " + str(len(self.ChiefT[4])) + "/" + str(Cert_CT))
            if len(self.MM[3]) < Qual_MM + Cert_MM:
                failure_reasons.append("  MM (Qualified): " + str(len(self.MM[3])) + "/" + str(Qual_MM + Cert_MM))
            if len(self.MM[4]) < Cert_MM:
                failure_reasons.append("  MM (Certified): " + str(len(self.MM[4])) + "/" + str(Cert_MM))
            if len(self.Clerk[3]) < Qual_Clerk + Cert_Clerk:
                failure_reasons.append(
                    "  Admin Desk (Qualified): " + str(len(self.Clerk[3])) + "/" + str(Qual_Clerk + Cert_Clerk)
                )
            if len(self.Clerk[4]) < Cert_Clerk:
                failure_reasons.append("  Admin Desk (Certified): " + str(len(self.Clerk[4])) + "/" + str(Cert_Clerk))
            if len(self.Starter[3]) < Qual_Starter + Cert_Starter:
                failure_reasons.append(
                    "  Starter (Qualified): " + str(len(self.Starter[3])) + "/" + str(Qual_Starter + Cert_Starter)
                )
            if len(self.Starter[4]) < Cert_Starter:
                failure_reasons.append("  Starter (Certified): " + str(len(self.Starter[4])) + "/" + str(Cert_Starter))
            if len(self.CFJ[3]) < Qual_CFJ + Cert_CFJ:
                failure_reasons.append("  CFJ (Qualified): " + str(len(self.CFJ[3])) + "/" + str(Qual_CFJ + Cert_CFJ))
            if len(self.CFJ[4]) < Cert_CFJ:
                failure_reasons.append("  CFJ (Certified): " + str(len(self.CFJ[4])) + "/" + str(Cert_CFJ))
            if len(self.IT[3]) < Qual_IT:
                failure_reasons.append("  IT (Qualified): " + str(len(self.IT[3])) + "/" + str(Qual_IT))
            if len(self.IT[4]) < Cert_IT:
                failure_reasons.append("  IT (Certified): " + str(len(self.IT[4])) + "/" + str(Cert_IT))
            if len(self.JoS[3]) < Qual_JoS:
                failure_reasons.append("  JoS (Qualified): " + str(len(self.JoS[3])) + "/" + str(Qual_JoS))
            if len(self.JoS[4]) < Cert_JoS:
                failure_reasons.append("  JoS (Certified): " + str(len(self.JoS[4])) + "/" + str(Cert_JoS))

            return failure_reasons

        # For efficiency, build the scenario from easiest to hardest positions to staff
        # This aids in early termination of the search.
        for x in range(Qual_CT):
            scenario["CT_Q" + str(x)] = self.ChiefT[3]
        for x in range(Cert_CT):
            scenario["CT_C" + str(x)] = self.ChiefT[4]
        for x in range(Qual_Clerk):
            scenario["Clerk_Q" + str(x)] = self.Clerk[3]
        for x in range(Cert_Clerk):
            scenario["Clerk_C" + str(x)] = self.Clerk[4]
        for x in range(Qual_Starter):
            scenario["Starter_Q" + str(x)] = self.Starter[3]
        for x in range(Cert_Starter):
            scenario["Starter_C" + str(x)] = self.Starter[4]
        for x in range(Qual_CFJ):
            scenario["CFJ_Q" + str(x)] = self.CFJ[3]
        for x in range(Cert_CFJ):
            scenario["CFJ_C" + str(x)] = self.CFJ[4]
        for x in range(Qual_MM):
            scenario["MM_Q" + str(x)] = self.MM[3]
        for x in range(Cert_MM):
            scenario["MM_C" + str(x)] = self.MM[4]
        for x in range(Level3):
            scenario["L3_" + str(x)] = []
            if self.Level_3_list:
                scenario["L3_" + str(x)].extend(self.Level_3_list)
            if self.Level_4_5s:
                scenario["L3_" + str(x)].extend(self.Level_4_5s)
        for x in range(Qual_Ref):
            scenario["L3Ref_" + str(x)] = []
            if self.Qualified_Refs:
                scenario["L3Ref_" + str(x)].extend([item[0] for item in self.Qualified_Refs])
            if self.Level_4_5s:
                scenario["L3Ref_" + str(x)].extend(self.Level_4_5s)
        for x in range(Level4_5):
            scenario["L45_" + str(x)] = self.Level_4_5s

        return scenario

    def _build_staffing_scenario_SandT(
        self, Qual_IT: int, Cert_IT: int, Qual_JoS: int, Cert_JoS: int, staff_list: list
    ) -> LIST_OR_DICT:
        """Build the dictionary for the stroke & turn scenario"""

        scenario = {}
        failure_reasons: list = []

        # Remove anyone already staffed on the senior grid

        qual_IT_left = list(filter(lambda i: i not in staff_list, self.IT[3]))
        cert_IT_left = list(filter(lambda i: i not in staff_list, self.IT[4]))
        qual_JoS_left = list(filter(lambda i: i not in staff_list, self.JoS[3]))
        cert_JoS_left = list(filter(lambda i: i not in staff_list, self.JoS[4]))

        # Check Quick Failure Conditions
        if (
            len(qual_IT_left) < Qual_IT + Cert_IT + Qual_JoS + Cert_JoS
            or len(cert_IT_left) < Cert_IT
            or len(qual_JoS_left) < Qual_JoS + Cert_JoS
            or len(cert_JoS_left) < Cert_JoS
        ):
            # Return each condition that failed as a list showing required and available
            if len(qual_IT_left) < Qual_IT + Cert_IT + Qual_JoS + Cert_JoS:
                failure_reasons.append(
                    " IT (Qualified): " + str(len(qual_IT_left)) + "/" + str(Qual_IT + Cert_IT + Qual_JoS + Cert_JoS)
                )
            if len(cert_IT_left) < Cert_IT:
                failure_reasons.append(" IT (Certified): " + str(len(cert_IT_left)) + "/" + str(Cert_IT))
            if len(qual_JoS_left) < Qual_JoS + Cert_JoS:
                failure_reasons.append(" JoS (Qualified): " + str(len(qual_JoS_left)) + "/" + str(Qual_JoS + Cert_JoS))
            if len(cert_JoS_left) < Cert_JoS:
                failure_reasons.append(" JoS (Certified): " + str(len(cert_JoS_left)) + "/" + str(Cert_JoS))

            return failure_reasons

        for x in range(Qual_IT):
            scenario["IT_Q" + str(x)] = qual_IT_left
        for x in range(Cert_IT):
            scenario["IT_C" + str(x)] = cert_IT_left
        for x in range(Qual_JoS):
            scenario["JoS_Q" + str(x)] = qual_JoS_left
        for x in range(Cert_JoS):
            scenario["JoS_C" + str(x)] = cert_JoS_left

        return scenario

    def _find_staffing_scenario(self, scenario: dict, current_plan: list, required_staff: int) -> list:
        """Try to find a workable staffing scenario"""
        working_plan: list = []

        if not scenario:  # No remaining jobs to staff
            return []

        # We need to make copies of the scenario and the current plan as we are going to check
        # This is recursive so we need to make sure we don't modify the original data

        scenario_copy = deepcopy(scenario)
        current_skill = scenario_copy.popitem()

        for name in current_skill[1]:
            working_plan = deepcopy(current_plan)
            if name not in current_plan and scenario_copy:
                working_plan = copy(current_plan)
                working_plan.append(name)
                if len(working_plan) > len(self.best_scenario):
                    self.best_scenario = copy(working_plan)
                sub = self._find_staffing_scenario(scenario_copy, working_plan, required_staff)
                if sub and len(sub) == required_staff:
                    if self.debug:
                        logging.debug(self.club_code + "Sub/Req: " + str(sub))
                    return sub
            elif name not in current_plan:
                if self.debug:
                    logging.debug(self.club_code + "Found: " + str(name))
                working_plan.append(name)
                if len(working_plan) > len(self.best_scenario):
                    self.best_scenario = copy(working_plan)
                return working_plan
        return []

    def dump_data_docx(self, doc: Document, club_fullname: str, reportdate: str, affiliates: list):
        """Produce the Word Document for the club"""
        doc.add_heading(club_fullname + " (" + self.club_code + ")", 0)
        doc.add_heading("Provisionally Approved Sanction Types (as of " + reportdate + ")", level=2)
        sanction_p = doc.add_paragraph()
        sanction_p.add_run("\n".join(self.Sanction_Level))
        if len(self.Sanction_Level) == 0:
            sanction_p.add_run("NO APPROVED SANCTION TYPES")

        doc.add_heading("Officials Summary", level=2)
        ps = doc.add_paragraph("No Level: %d\n" % self.Level_None)
        ps.add_run("Level 1 : %d\n" % self.Level_1s)
        ps.add_run("Level 2 : %d\n" % self.Level_2s)
        ps.add_run("Level 3 : %d\n" % self.Level_3s)
        ps.add_run("Level 4 : %d\n" % self.Level_4s)
        ps.add_run("Level 5 : %d" % self.Level_5s)

        doc.add_heading("Skills Summary (Excludes Level 4/5)", level=3)
        table = doc.add_table(rows=1, cols=4)
        row = table.rows[0].cells
        row[0].text = "Qualfication"
        row[1].text = "Total Clinics"
        row[2].text = "1 Sign-Off"
        row[3].text = "2 Sign-Offs"
        row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        table_data = [
            ("Intro to Swimming", str(self.Intro[0]), str(self.Intro[1]), str(self.Intro[2])),
            ("Stroke & Turn (Pre Sept/23)", str(self.SandT[0]), str(self.SandT[1]), str(self.SandT[2])),
            ("Inspector of Turns", str(self.IT[0]), str(self.IT[1]), str(self.IT[2])),
            ("Judge of Stroke", str(self.JoS[0]), str(self.JoS[1]), str(self.JoS[2])),
            ("Chief Timekeeper", str(self.ChiefT[0]), str(self.ChiefT[1]), str(self.ChiefT[2])),
            ("Admin Desk (Clerk)", str(self.Clerk[0]), str(self.Clerk[1]), str(self.Clerk[2])),
            ("Meet Manager", str(self.MM[0]), str(self.MM[1]), str(self.MM[2])),
            ("Starter", str(self.Starter[0]), str(self.Starter[1]), str(self.Starter[2])),
            ("CFJ/CJE", str(self.CFJ[0]), str(self.CFJ[1]), str(self.CFJ[2])),
            ("Chief Recorder/Recorder", str(self.RecSec[0]), "", ""),
            ("Referee", str(self.Referee[0]), "", ""),
        ]

        for entry in table_data:
            row = table.add_row().cells
            for k in range(len(entry)):
                row[k].text = entry[k]
                row[k].width = Inches(2.5) if k == 0 else Inches(1.5)
                row[k].paragraphs[0].alignment = (
                    docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER if k > 0 else docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                )

        table.style = "Light Grid Accent 5"

        if len(self.Qualified_Refs) > 0:
            doc.add_heading("Qualified Level III Referees", level=3)
            refp = doc.add_paragraph()
            refp.add_run(
                "\n".join(
                    ref[0]
                    + (" (Para eModule)" if ref[2] == "yes" else "")
                    + (" (Para Domestic: " + ref[1] + ")" if not pd.isnull(ref[1]) else "")
                    for ref in self.Qualified_Refs
                )
            )

        if self._config.get_bool("incl_affiliates") and affiliates:
            affiliated_officials = self._club_data_full[self._club_data_full["Registration Id"].isin(affiliates)]
            if not affiliated_officials.empty:  # We could have affiliated officials but no supporting data
                doc.add_heading("Affiliated Officials", level=3)

                atable = doc.add_table(rows=1, cols=3)
                row = atable.rows[0].cells
                row[0].text = "Registration Id"
                row[1].text = "Name"
                row[2].text = "Certification Level"
                row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                for index, affiliated_official in affiliated_officials.iterrows():
                    row = atable.add_row().cells
                    row[0].text = affiliated_official["Registration Id"]
                    row[1].text = (
                        affiliated_official["Last Name"]
                        + ", "
                        + affiliated_official["First Name"]
                        + " ("
                        + affiliated_official["ClubCode"]
                        + ")"
                    )
                    row[2].text = affiliated_official["Current_CertificationLevel"]
                    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                atable.style = "Light Grid Accent 5"
                atable.allow_autofit = True

        if self._config.get_bool("incl_sanction_errors") and self.Failed_Sanctions:
            doc.add_heading("Sanctioning Issues", level=3)
            error_p = doc.add_paragraph()
            error_p.add_run("\n".join(self.Failed_Sanctions))

        if self._config.get_bool("incl_errors"):
            if self.NoLevel_Missing_Cert:
                doc.add_heading("RTR Error - Official(s) missing Level I Certification Record", level=2)
                error_p1 = doc.add_paragraph()
                error_p1.add_run("\n".join(self.NoLevel_Missing_Cert))

            if self.Missing_Level_II:
                doc.add_heading("RTR Error - Official(s) missing Level II Certification Record", level=2)
                error_p3 = doc.add_paragraph()
                error_p3.add_run("\n".join(self.Missing_Level_II))

            if self.Missing_Level_III:
                doc.add_heading("RTR Possible Error - Official(s) missing Level III Certification Record", level=2)
                error_p2 = doc.add_paragraph()
                error_p2.add_run("* COA to check last certification date and para certification status\n")
                error_p2.add_run("\n".join(self.Missing_Level_III))

            if self.NoLevel_Missing_SM:
                doc.add_heading("RTR Warning - Level I Partially Complete - Need Safety Marshal", level=2)
                warn_p = doc.add_paragraph()
                warn_p.add_run("\n".join(self.NoLevel_Missing_SM))

            if self.NoLevel_Has_II:
                doc.add_heading("RTR Warning - Official has Level II clinics - Missing Level I Certification", level=2)
                warn_p2 = doc.add_paragraph()
                warn_p2.add_run("\n".join(self.NoLevel_Has_II))
