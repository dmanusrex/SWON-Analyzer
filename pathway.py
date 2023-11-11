# Club Analyzer - https://github.com/dmanusrex/SWON-Analyzer
# Copyright (C) 2023 - Darren Richer
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.


"""  Experimental Module - Generate new pathway migration documents """

import logging
import os
from datetime import datetime
from threading import Thread
from tkinter import BooleanVar, StringVar, filedialog


import customtkinter as ctk  # type: ignore
import docx  # type: ignore
import pandas as pd
from docx import Document  # type: ignore
from docx.shared import Inches  # type: ignore
from docxcompose.composer import Composer  # type: ignore
from slugify import slugify

# Appliction Specific Imports
from config import AnalyzerConfig
from CTkMessagebox import CTkMessagebox  # type: ignore
from rtr import RTR
from tooltip import ToolTip
from ui_common import Officials_Status_Frame


class Pathway_Documents_Frame(ctk.CTkFrame):
    """Generate Word Documents from a supplied RTR file"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig, rtr: RTR):
        super().__init__(container)
        self._config = config
        self._rtr = rtr

        # Get the needed options
        self._officials_list = StringVar(value=self._config.get_str("officials_list"))
        self._officials_list_filename = StringVar(value=os.path.basename(self._officials_list.get()))
        self._np_report_directory = StringVar(value=self._config.get_str("np_report_directory"))
        self._np_report_file = StringVar(value=self._config.get_str("np_report_file_docx"))

        # Add support for the club selection
        self._club_list = ["None"]
        self._club_selected = ctk.StringVar(value="None")

        # self is a vertical container that will contain 3 frames
        self.columnconfigure(0, weight=1)

        # Setup the sub-frames
        ctk.CTkLabel(self, text="New Pathway Recommendations").grid(column=0, row=0, sticky="we", pady=(10, 0))

        club_select_frame = ctk.CTkFrame(self)
        club_select_frame.grid(column=0, row=1, columnspan=2, sticky="news", padx=10, pady=10)

        Officials_Status_Frame(self, self._config).grid(column=0, row=2, sticky="news")

        filesframe = ctk.CTkFrame(self)
        filesframe.grid(column=0, row=3, sticky="news", padx=10, pady=10)
        filesframe.rowconfigure(0, weight=1)

        buttonsframe = ctk.CTkFrame(self)
        buttonsframe.grid(column=0, row=4, sticky="news", padx=10, pady=10)
        buttonsframe.rowconfigure(0, weight=0)

        # Club selection

        self.club_dropdown = ctk.CTkOptionMenu(
            club_select_frame, dynamic_resizing=True, values=self._club_list, variable=self._club_selected
        )
        self.club_dropdown.grid(row=0, column=1, padx=20, pady=(10, 10), sticky="w")
        ctk.CTkLabel(club_select_frame, text="Club", anchor="w").grid(
            row=0, column=0, sticky="w", padx=10, pady=(10, 10)
        )

        # Files Section
        ctk.CTkLabel(filesframe, text="Files and Directories").grid(column=0, row=0, sticky="w", padx=10, columnspan=2)

        btn2 = ctk.CTkButton(filesframe, text="Pathway Docs Folder", command=self._handle_report_dir_browse)
        btn2.grid(column=0, row=1, padx=10, pady=10, sticky="ew")
        ToolTip(btn2, text="Select where output files will be sent")
        ctk.CTkLabel(filesframe, textvariable=self._np_report_directory).grid(
            column=1, row=1, sticky="w", padx=(0, 10)
        )

        btn3 = ctk.CTkButton(filesframe, text="Consolidated Report File", command=self._handle_report_file_browse)
        btn3.grid(column=0, row=2, padx=10, pady=10, sticky="ew")
        ToolTip(btn3, text="Set report file name")
        ctk.CTkLabel(filesframe, textvariable=self._np_report_file).grid(column=1, row=2, sticky="w", padx=(0, 10))

        # Add Command Button

        ctk.CTkLabel(buttonsframe, text="Actions").grid(column=0, row=0, sticky="w", padx=10)

        self.reports_btn = ctk.CTkButton(buttonsframe, text="Generate Reports", command=self._handle_reports_btn)
        self.reports_btn.grid(column=0, row=1, sticky="news", padx=10, pady=10)

        self.bar = ctk.CTkProgressBar(master=buttonsframe, orientation="horizontal", mode="indeterminate")

        # Register Callback
        self._rtr.register_update_callback(self.refresh_club_list)

    def refresh_club_list(self):
        self._club_list = ["None"]
        self._club_list = self._club_list + [club[1] for club in self._rtr.club_list_names]

        self.club_dropdown.configure(values=self._club_list)
        if len(self._club_list) == 2:  # There is only one club loaded, change default
            self.club_dropdown.set(self._club_list[1])
        else:
            self.club_dropdown.set(self._club_list[0])
        logging.info("New Pathway Module - Club List Refreshed")

    def _handle_report_dir_browse(self) -> None:
        directory = filedialog.askdirectory()
        if len(directory) == 0:
            return
        directory = os.path.normpath(directory)
        self._config.set_str("np_report_directory", directory)
        self._np_report_directory.set(directory)

    def _handle_report_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            title="Report File",
            initialfile=os.path.basename(self._np_report_file.get()),
            initialdir=self._config.get_str("np_report_directory"),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("np_report_file_docx", report_file)
        self._np_report_file.set(report_file)

    def buttons(self, newstate) -> None:
        """Enable/disable all buttons"""
        self.reports_btn.configure(state=newstate)

    def _handle_reports_btn(self) -> None:
        if self._rtr.rtr_data.empty:
            logging.info("Load data first...")
            CTkMessagebox(master=self, title="Error", message="Load RTR Data First", icon="cancel", corner_radius=0)
            return
        club = self._club_selected.get()
        if club == "None":
            logging.info("Select a club first...")
            CTkMessagebox(master=self, title="Error", message="Select a club first", icon="cancel", corner_radius=0)
            return
        self.buttons("disabled")
        self.bar.grid(row=2, column=0, pady=10, padx=20, sticky="s")
        self.bar.set(0)
        self.bar.start()
        reports_thread = Generate_Reports(self._rtr, self._config, club)
        reports_thread.start()
        self.monitor_reports_thread(reports_thread)

    def monitor_reports_thread(self, thread):
        if thread.is_alive():
            # check the thread every 100ms
            self.after(100, lambda: self.monitor_reports_thread(thread))
        else:
            self.buttons("enabled")
            self.bar.stop()
            self.bar.grid_forget()
            thread.join()


class Pathway_ROR_Frame(ctk.CTkFrame):
    """Reports Setting and Generation for RORs and POAs"""

    def __init__(self, container: ctk.CTk, config: AnalyzerConfig, rtr: RTR):
        super().__init__(container)  # , text="General Settings", padding=5
        self._config = config
        self._rtr = rtr
        self._report_file = StringVar(value=self._config.get_str("np_ror_file_docx"))
        self._report_csv = StringVar(value=self._config.get_str("np_report_file_csv"))

        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)

        # Setup the sub-frames
        ctk.CTkLabel(self, text="ROR/POA New Pathway Reports (Multi-Club)").grid(column=0, row=0, sticky="we", pady=10)

        optionsframe = ctk.CTkFrame(self)
        optionsframe.grid(column=0, row=2, sticky="news", padx=10, pady=10)

        Officials_Status_Frame(self, self._config).grid(column=0, row=1, sticky="news")

        filesframe = ctk.CTkFrame(self)
        filesframe.grid(column=0, row=4, sticky="news", padx=10, pady=10)
        filesframe.rowconfigure(0, weight=1)
        filesframe.rowconfigure(1, weight=1)
        filesframe.rowconfigure(2, weight=1)

        buttonsframe = ctk.CTkFrame(self)
        buttonsframe.grid(column=0, row=6, sticky="news", padx=10, pady=10)
        buttonsframe.rowconfigure(0, weight=1)
        buttonsframe.columnconfigure(0, weight=1)

        # Report Options

        self._gen_np_csv_var = BooleanVar(optionsframe, value=self._config.get_bool("gen_np_csv"))
        ctk.CTkSwitch(
            optionsframe,
            text="Create CSV File",
            variable=self._gen_np_csv_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_gen_csv_file,
        ).grid(column=1, row=0, sticky="news", padx=20, pady=10)

        self._gen_np_warnings_var = BooleanVar(optionsframe, value=self._config.get_bool("gen_np_warnings"))
        ctk.CTkSwitch(
            optionsframe,
            text="Generate Warnings Report",
            variable=self._gen_np_warnings_var,
            onvalue=True,
            offvalue=False,
            command=self._handle_gen_warnings,
        ).grid(column=1, row=2, sticky="news", padx=20, pady=10)

        # Report File

        rptbtn = ctk.CTkButton(filesframe, text="Warnings Filename", command=self._handle_report_file_browse)
        rptbtn.grid(column=0, row=0, padx=20, pady=10)
        ToolTip(rptbtn, text="The name of the warnings/exceptions file")
        ctk.CTkLabel(filesframe, textvariable=self._report_file).grid(column=1, row=0, sticky="w", padx=(0, 10))

        csvbtn = ctk.CTkButton(filesframe, text="CSV Filename", command=self._handle_csv_browse)
        csvbtn.grid(column=0, row=2, padx=20, pady=10)
        ToolTip(csvbtn, text="The name of the consolidated CSV file")
        ctk.CTkLabel(filesframe, textvariable=self._report_csv).grid(column=1, row=2, sticky="w", padx=(0, 10))

        # Action Button

        self.reports_btn = ctk.CTkButton(buttonsframe, text="Generate Reports", command=self._handle_reports_btn)
        self.reports_btn.grid(column=0, row=0, sticky="ew", padx=20, pady=10)

        self.bar = ctk.CTkProgressBar(master=buttonsframe, orientation="horizontal", mode="indeterminate")

    def buttons(self, newstate) -> None:
        """Enable/disable all buttons on the UI"""
        self.reports_btn.configure(state=newstate)

    def _handle_gen_csv_file(self, *_arg) -> None:
        self._config.set_bool("gen_np_csv", self._gen_np_csv_var.get())

    def _handle_gen_warnings(self, *_arg) -> None:
        self._config.set_bool("gen_np_warnings", self._gen_np_warnings_var.get())

    def _handle_report_file_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            title="Warnings/Exceptions File",
            initialfile=os.path.basename(self._report_file.get()),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("np_ror_file_docx", report_file)
        self._report_file.set(report_file)

    def _handle_csv_browse(self) -> None:
        report_file = filedialog.asksaveasfilename(
            filetypes=[("CSV Files", "*.csv")],
            defaultextension=".csv",
            title="CSV Filename",
            initialfile=os.path.basename(self._report_file.get()),
        )
        if len(report_file) == 0:
            return
        self._config.set_str("np_report_file_csv", report_file)
        self._report_file.set(report_file)

    def _handle_reports_btn(self) -> None:
        if self._rtr.rtr_data.empty:
            logging.info("Load data first...")
            CTkMessagebox(master=self, title="Error", message="Load RTR Data First", icon="cancel", corner_radius=0)
            return
        self.buttons("disabled")
        self.bar.grid(row=2, column=0, pady=10, padx=20, sticky="s")
        self.bar.set(0)
        self.bar.start()
        reports_thread = _Generate_NP_ROR_Reports(self._rtr, self._config)
        reports_thread.start()
        self.monitor_reports_thread(reports_thread)

    def monitor_reports_thread(self, thread) -> None:
        if thread.is_alive():
            # check the thread every 100ms
            self.after(100, lambda: self.monitor_reports_thread(thread))
        else:
            self.bar.stop()
            self.bar.grid_forget()
            self.buttons("enabled")
            thread.join()


class _Generate_NP_ROR_Reports(Thread):
    def __init__(self, rtr: RTR, config: AnalyzerConfig):
        super().__init__()
        self._rtr: pd.DataFrame = rtr.rtr_data
        self._config: AnalyzerConfig = config

    def run(self):
        logging.info("Reporting in Progress...")

        self._report_file = self._config.get_str("np_ror_file_docx")
        self._report_csv = self._config.get_str("np_report_file_csv")

        _gen_np_warnings = self._config.get_bool("gen_np_warnings")
        _gen_np_csv = self._config.get_bool("gen_np_csv")

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        self._rtr_filtered = self._rtr[self._rtr["Status"].isin(status_values)]

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        if _gen_np_csv:
            logging.info("Generating CSV Report")
            key_columns = [
                "Registration Id",
                "First Name",
                "Last Name",
                "Club",
                "Region",
                "Province",
                "Status",
                "Current_CertificationLevel",
                "Intro_Status",
                "ST_Status",
                "IT_Status",
                "JoS_Status",
                "CT_Status",
                "Admin_Status",
                "MM_Status",
                "Starter_Status",
                "CFJ_Status",
                "ChiefRec_Status",
                "Referee_Status",
                "Para Swimming eModule",
                "NP_Official",
                "NP_Ref1",
                "NP_Ref2",
                "NP_Starter1",
                "NP_Starter2",
                "NP_MM1",
                "NP_MM2",
            ]
            try:
                self._rtr_filtered.to_csv(self._report_csv, columns=key_columns, index=False)
            except Exception as e:
                logging.info("Unable to save CSV file: {}".format(type(e).__name__))
                logging.info("Exception message: {}".format(e))
                CTkMessagebox(title="Error", message="Unable to save CSV file", icon="cancel", corner_radius=0)

            logging.info("CSV Report Complete")

        if _gen_np_warnings:
            doc = Document()
            logging.info("Generating Warnings Report")

            #           Place holder for warnings report
            doc.add_heading("Warnings Report", 0)
            doc.add_paragraph("Report Generated: " + report_time)

        try:
            doc.save(self._report_file)
        except Exception as e:
            logging.info("Unable to save full report: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))
            CTkMessagebox(title="Error", message="Unable to save full report file", icon="cancel", corner_radius=0)

        CTkMessagebox(title="Reports", message="Reports complete", icon="check", option_1="OK", corner_radius=0)

        logging.info("Reports Complete")


class NewPathway:
    def __init__(self, club: str, club_data_set: pd.DataFrame, config: AnalyzerConfig, **kwargs):
        self._club_data = club_data_set.copy()
        self.club_code = club
        self._config = config

    def _is_valid_date(self, date_string) -> bool:
        if pd.isnull(date_string):
            return False
        if date_string == "0001-01-01":
            return False
        try:
            datetime.strptime(date_string, "%Y-%m-%d")
            return True
        except ValueError:
            return False

    def _get_date(self, date_string) -> str:
        if pd.isnull(date_string):
            return ""
        if date_string == "0001-01-01":
            return ""
        return date_string

    def _count_signoffs(self, clinic_date_1, clinic_date_2) -> int:
        count = 0
        if self._is_valid_date(clinic_date_1):
            count += 1
        if self._is_valid_date(clinic_date_2):
            count += 1
        return count

    def add_clinic(self, table, clinic_name, clinic_date, signoff_1, signoff_2) -> None:
        row = table.add_row().cells
        row[0].text = clinic_name
        row[1].text = self._get_date(clinic_date)
        row[2].text = self._get_date(signoff_1)
        row[3].text = self._get_date(signoff_2)
        row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTERAdmin_Status

    def add_pathway(self, table, pathway_progression, certified) -> None:
        row = table.add_row().cells
        row[0].text = pathway_progression
        row[1].text = certified
        row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def export_csv(self, filename: str) -> None:
        """Export the club data to a CSV file"""

        key_columns = [
            "Registration Id",
            "First Name",
            "Last Name",
            "Club",
            "Region",
            "Province",
            "Status",
            "Current_CertificationLevel",
            "Intro_Status",
            "ST_Status",
            "IT_Status",
            "JoS_Status",
            "CT_Status",
            "Admin_Status",
            "MM_Status",
            "Starter_Status",
            "CFJ_Status",
            "ChiefRec_Status",
            "Referee_Status",
            "Para Swimming eModule",
            "NP_Official",
            "NP_Ref1",
            "NP_Ref2",
            "NP_Starter1",
            "NP_Starter2",
            "NP_MM1",
            "NP_MM2",
        ]
        try:
            self._club_data.to_csv(filename, columns=key_columns, index=False)
        except Exception as e:
            logging.info("Unable to save CSV file: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))
            CTkMessagebox(title="Error", message="Unable to save CSV file", icon="cancel", corner_radius=0)

    def dump_data_docx(self, club_fullname: str, reportdate: str) -> list:
        """Produce the Word Document for the club and return a list of files"""

        _report_directory = self._config.get_str("np_report_directory")
        _email_list_csv = self._config.get_str("email_list_csv")
        csv_list = []  # CSV entries for email list (Lastname, Firstname, E-Mail address and Filename)

        for index, entry in self._club_data.iterrows():
            # create a filename from the last and firstnames using slugify and the report directory

            filename = os.path.abspath(
                os.path.join(_report_directory, slugify(entry["Last Name"] + "_" + entry["First Name"]) + ".docx")
            )
            csv_list.append([entry["Last Name"], entry["First Name"], entry["Email"], filename])

            doc = Document()

            doc.add_heading("2023/24 New Pathway Mapping", 0)

            p = doc.add_paragraph()
            p.add_run("Report Date: " + reportdate)
            p.add_run(
                "\n\nName: "
                + entry["Last Name"]
                + ", "
                + entry["First Name"]
                + " (SNC ID # "
                + entry["Registration Id"]
                + ")"
            )
            p.add_run("\n\nClub: " + club_fullname + " (" + self.club_code + ")")
            p.add_run("\n\nCurrent Certification Level: ")
            p.add_run(
                "NONE" if pd.isnull(entry["Current_CertificationLevel"]) else entry["Current_CertificationLevel"]
            )

            table = doc.add_table(rows=1, cols=4)
            row = table.rows[0].cells
            row[0].text = "Clinic"
            row[1].text = "Clinic Date"
            row[2].text = "Sign Off #1"
            row[3].text = "Sign Off #2"
            row[0].width = Inches(2.0)
            row[1].width = Inches(1.5)
            row[2].width = Inches(1.5)
            row[3].width = Inches(1.5)
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            self.add_clinic(
                table,
                "Intro to Swimming",
                entry["Introduction to Swimming Officiating-ClinicDate"],
                entry["Introduction to Swimming Officiating-Deck Evaluation #1 Date"],
                entry["Introduction to Swimming Officiating-Deck Evaluation #2 Date"],
            )
            self.add_clinic(table, "Safety Marshal", entry["Safety Marshal-ClinicDate"], "N/A", "N/A")
            self.add_clinic(
                table,
                "Stroke & Turn (Pre Sept/23)",
                entry["Judge of Stroke/Inspector of Turns-ClinicDate"],
                entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #1 Date"],
                entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Inspector of Turns",
                entry["Inspector of Turns-ClinicDate"],
                entry["Inspector of Turns-Deck Evaluation #1 Date"],
                entry["Inspector of Turns-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Judge of Stroke",
                entry["Judge of Stroke-ClinicDate"],
                entry["Judge of Stroke-Deck Evaluation #1 Date"],
                entry["Judge of Stroke-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Chief Timekeeper",
                entry["Chief Timekeeper-ClinicDate"],
                entry["Chief Timekeeper-Deck Evaluation #1 Date"],
                entry["Chief Timekeeper-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Admin Desk (Clerk)",
                entry["Administration Desk (formerly Clerk of Course) Clinic-ClinicDate"],
                entry["Administration Desk (formerly Clerk of Course) Clinic-Deck Evaluation #1 Date"],
                entry["Administration Desk (formerly Clerk of Course) Clinic-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Meet Manager",
                entry["Meet Manager-ClinicDate"],
                entry["Meet Manager-Deck Evaluation #1 Date"],
                entry["Meet Manager-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Starter",
                entry["Starter-ClinicDate"],
                entry["Starter-Deck Evaluation #1 Date"],
                entry["Starter-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "CFJ/CJE",
                entry["Chief Finish Judge/Chief Judge-ClinicDate"],
                entry["Chief Finish Judge/Chief Judge-Deck Evaluation #1 Date"],
                entry["Chief Finish Judge/Chief Judge-Deck Evaluation #2 Date"],
            )
            self.add_clinic(
                table,
                "Chief Recorder/Recorder",
                entry["Chief Recorder and Recorder (formerly Recorder/Scorer) Clinic-ClinicDate"],
                "N/A",
                "N/A",
            )
            self.add_clinic(table, "Referee", entry["Referee-ClinicDate"], "N/A", "N/A")
            self.add_clinic(table, "Para eModule", entry["Para Swimming eModule-ClinicDate"], "N/A", "N/A")

            table.style = "Light Grid Accent 5"
            table.autofit = True

            # Add logic to define pathway progression

            doc.add_heading("New Pathway Progression", 2)
            nptable = doc.add_table(rows=1, cols=4)
            row = nptable.rows[0].cells
            row[0].text = "New Pathway Level"
            row[1].text = "Certifed?"
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            self.add_pathway(nptable, "Certified Official", entry["NP_Official"])
            self.add_pathway(nptable, "Referee 1", entry["NP_Ref1"])
            self.add_pathway(nptable, "Referee 2", entry["NP_Ref2"])
            self.add_pathway(nptable, "Starter 1", entry["NP_Starter1"])
            self.add_pathway(nptable, "Starter 2", entry["NP_Starter2"])
            self.add_pathway(nptable, "Meet Manager 1", entry["NP_MM1"])
            self.add_pathway(nptable, "Meet Manager 2", entry["NP_MM2"])

            nptable.style = "Light Grid Accent 5"
            nptable.autofit = True

            # Recommendations to be added here

            try:
                doc.save(filename)

            except Exception as e:
                logging.info(
                    f'Error processing offiical {entry["Last Name"]}, {entry["First Name"]}: {type(e).__name__} - {e}'
                )

        return csv_list


class Generate_Reports(Thread):
    def __init__(self, rtr: RTR, config: AnalyzerConfig, selected_club: str):
        super().__init__()
        self._rtr = rtr
        self._df: pd.DataFrame = self._rtr.rtr_data
        self._config: AnalyzerConfig = config
        self._selected_club = selected_club

    def run(self):
        # This still has code elements to run multiple clubs. For single clubs a simple filter has been applied.
        # Once the final specifications are set, this can be greatly simplified.

        logging.info("Reporting in Progress...")

        _report_directory = self._config.get_str("np_report_directory")
        _report_file_docx = self._config.get_str("np_report_file_docx")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_docx))
        _email_list_csv = self._config.get_str("email_list_csv")
        _full_csv_file = os.path.abspath(os.path.join(_report_directory, _email_list_csv))

        club_list_names = self._rtr.club_list_names

        club_summaries = []

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        all_csv_entries = []

        for club, club_full in filter(lambda x: x[1] == self._selected_club, club_list_names):
            logging.info("Processing %s" % club_full)
            club_data = self._df[(self._df["ClubCode"] == club)]
            club_data = club_data[club_data["Status"].isin(status_values)]
            club_stat = NewPathway(club, club_data, self._config)
            club_csv = club_stat.dump_data_docx(club_full, report_time)
            all_csv_entries.extend(club_csv)
            club_summaries.append([club, club_full, club_stat])

        # Create the email list CSV file
        #
        # The email list is a CSV file with the following columns:
        #  Last Name, First Name, E-Mail address, Filename

        logging.info("Creating email list CSV file")
        email_list_df = pd.DataFrame(all_csv_entries, columns=["Last Name", "First Name", "EMail", "Filename"])

        try:
            email_list_df.to_csv(_full_csv_file, index=False)
        except Exception as e:
            logging.info("Unable to save email list: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))

        # Create the master document

        logging.info("Creating master document")

        number_of_sections = len(all_csv_entries)
        master = Document()
        composer = Composer(master)
        for i in range(0, number_of_sections):
            doc_temp = Document(all_csv_entries[i][3])
            doc_temp.add_page_break()
            composer.append(doc_temp)

        try:
            composer.save(_full_report_file)
        except Exception as e:
            logging.info("Unable to save full report: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))

        logging.info("Report Complete")


def main():
    """testing"""
    root = ctk.CTk()
    root.resizable(True, True)
    options = AnalyzerConfig()
    rtrdata = RTR(options)
    #    settings = Pathway_Documents_Frame(root, options, rtrdata)
    settings = Pathway_ROR_Frame(root, options, rtrdata)
    settings.grid(column=0, row=0, sticky="news")
    root.mainloop()


if __name__ == "__main__":
    main()
